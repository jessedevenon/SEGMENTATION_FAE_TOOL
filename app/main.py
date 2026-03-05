"""
CRM Segmentation Clientèle FAE 2026 - Version 7.3 Premium
Outil d'analyse et de priorisation de portefeuille clients
Cabinet Expert-Comptable - Réforme Facturation Électronique
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import io
import hmac

from analyzer import FAEAnalyzer
from style import inject_custom_css, render_premium_header, metric_card_html, section_divider

# Vérification disponibilité python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ========================================
# PROTECTION PAR MOT DE PASSE
# ========================================

def check_password():
    """Vérifie si l'utilisateur a entré le bon mot de passe."""

    def password_entered():
        """Vérifie si le mot de passe saisi est correct."""
        if hmac.compare_digest(st.session_state["password"], "Compta07!"):
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Ne pas stocker le mot de passe
        else:
            st.session_state["password_correct"] = False

    # Si déjà validé, retourner True
    if st.session_state.get("password_correct", False):
        return True

    # Afficher l'écran de connexion
    st.markdown("""
    <div style="
        max-width: 400px;
        margin: 10rem auto;
        padding: 3rem;
        background: linear-gradient(135deg, #dbeafe 0%, #bfdbfe 100%);
        border: 2px solid #60a5fa;
        border-radius: 20px;
        text-align: center;
        box-shadow: 0 8px 32px rgba(96, 165, 250, 0.3);
    ">
        <h1 style="color: #000000 !important; margin-bottom: 1rem; font-weight: 700;">🔒 Accès Sécurisé</h1>
        <p style="color: #1e293b !important; margin-bottom: 2rem; font-weight: 500; line-height: 1.6;">
            Outil d'Analyse & Segmentation Client RFE<br>
            Réservé aux cabinets partenaires
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.text_input(
            "🔑 Mot de passe",
            type="password",
            on_change=password_entered,
            key="password",
            placeholder="Entrez le mot de passe cabinet"
        )
        
        if "password_correct" in st.session_state:
            st.error("❌ Mot de passe incorrect. Contactez votre administrateur.")
    
    st.stop()  # Arrête l'exécution si pas de mot de passe


# Vérification du mot de passe AVANT tout le reste
if not check_password():
    st.stop()

# ========================================
# FIN PROTECTION - DÉBUT APP NORMALE
# ========================================

# Configuration page
st.set_page_config(
    page_title="Outil d'Analyse & Segmentation Client RFE",
    page_icon="🏆",
    layout="wide"
)

# Le reste de votre code continue ici...

# ========================================
# CONFIGURATION PAGE
# ========================================
st.set_page_config(
    page_title="Outil d'Analyse & Segmentation Client RFE",
    page_icon="🏆",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================================
# INJECTION CSS PREMIUM
# ========================================
st.markdown(inject_custom_css(), unsafe_allow_html=True)

# Forcer le masquage du bouton "réduire la sidebar" : CSS + JS en secours
st.markdown("""
<style>
html body [data-testid="collapsedControl"],
html body [data-testid*="collapsed"],
html body [data-testid*="Collapse"],
html body button[aria-label*="collapse"],
html body button[aria-label*="Collapse"],
html body button[aria-label*="sidebar"],
html body button[aria-label*="Sidebar"],
html body [class*="sidebarCollapse"],
html body [class*="sidebar-collapse"] {
    display: none !important;
    visibility: hidden !important;
    pointer-events: none !important;
    width: 0 !important;
    height: 0 !important;
    min-width: 0 !important;
    min-height: 0 !important;
    opacity: 0 !important;
    position: absolute !important;
    left: -9999px !important;
    z-index: -9999 !important;
}
</style>
<script>
(function() {
  function hideSidebarCollapse() {
    var nodes = [];
    document.querySelectorAll('[data-testid="collapsedControl"],[data-testid*="collapsed"],[data-testid*="Collapse"]').forEach(function(el) { nodes.push(el); });
    ['collapse','Collapse','sidebar','Sidebar','réduire'].forEach(function(word) {
      document.querySelectorAll('button[aria-label*="' + word + '"]').forEach(function(b) { nodes.push(b); });
    });
    nodes.forEach(function(n) {
      if (!n || !n.style) return;
      n.style.setProperty('display', 'none', 'important');
      n.style.setProperty('visibility', 'hidden', 'important');
      n.style.setProperty('pointer-events', 'none', 'important');
      n.style.setProperty('width', '0', 'important');
      n.style.setProperty('height', '0', 'important');
      n.style.setProperty('opacity', '0', 'important');
      n.style.setProperty('position', 'absolute', 'important');
      n.style.setProperty('left', '-9999px', 'important');
      n.setAttribute('tabindex', '-1');
      n.setAttribute('aria-hidden', 'true');
      if (n.tagName === 'BUTTON') n.disabled = true;
    });
    var sidebar = document.querySelector('section[data-testid="stSidebar"]');
    if (sidebar && sidebar.nextElementSibling) {
      var sib = sidebar.nextElementSibling;
      if (String(sib.getAttribute('data-testid') || '').indexOf('collapse') !== -1) {
        sib.style.setProperty('display', 'none', 'important');
        sib.style.setProperty('pointer-events', 'none', 'important');
      }
    }
  }
  hideSidebarCollapse();
  setInterval(hideSidebarCollapse, 250);
  var obs = new MutationObserver(hideSidebarCollapse);
  if (document.body) obs.observe(document.body, { childList: true, subtree: true });
})();
</script>
""", unsafe_allow_html=True)

# ========================================
# SIDEBAR PREMIUM (toujours affichée, pas de réduction)
# ========================================
with st.sidebar:
    # En-tête compact (icône + court libellé)
    st.markdown("""
    <div class="sidebar-logo">
        <span style="font-size: 1.9rem;">🏆</span>
        <span class="sidebar-title">Réforme de la facturation électronique</span>
    </div>
    """, unsafe_allow_html=True)
    
    # Navigation
    page = st.radio(
        "Navigation",
        [
            "🏠 Accueil",
            "📊 Dashboard",
            "🎯 Guide Missions",
            "🔍 Matrice 2.0",
            "💰 Simulateur CA",
            "📅 Plan d'Action",
            "⚠️ Analyse Risques",
            "📚 Bibliothèque",
            "📄 Livrables Word",
            "🔧 Budget TCO",
            "📤 Exports"
        ],
        label_visibility="collapsed"
    )
    
    # Footer sidebar compact
    st.markdown("""
    <div class="sidebar-footer">
        <p>🔒 Local · Confidentiel</p>
    </div>
    """, unsafe_allow_html=True)

# ========================================
# PARAMETRES & CONFIGURATION
# ========================================

# Coefficients scoring
coefficients = {
    "outil_non": 10,
    "outil_part": 5,
    "outil_oui": 1
}

# Tarifs missions
tarifs = {
    "audit_min": 1200,
    "audit_max": 1500,
    "formation_min": 600,
    "formation_max": 800,
    "info_min": 150,
    "info_max": 300
}

# Secteurs
SECTEURS_VALIDES = [
    "Autres",
    "Commerce de détail",
    "Hôtellerie",
    "Réparation de véhicules",
    "Activités juridiques et comptables",
    "Immobilier",
    "Industrie et fabrication",
    "Activités financières",
    "Autres services aux personnes",
    "Santé",
    "Travaux de construction",
    "Agriculture",
    "Architecture et ingénierie",
    "Enseignement",
    "Construction de bâtiments",
    "Transport",
    "Sports et loisirs",
    "Activités informatiques",
    "Services de conseil aux entreprises",
    "Industrie agroalimentaire",
    "Production audiovisuelle",
    "Édition",
    "Maintenance",
]

# ========================================
# FONCTION : CRÉATION TEMPLATE EXCEL (1 feuille uniquement)
# ========================================
def creer_template_excel():
    """Crée le modèle Excel avec une seule feuille Données (2-3 lignes d'exemples)."""
    import xlsxwriter
    from io import BytesIO
    
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
    # === UNE SEULE FEUILLE : DONNÉES ===
    worksheet_data = workbook.add_worksheet('Données')
    
    # Formats
    header_format = workbook.add_format({
        'bold': True,
        'font_color': 'white',
        'bg_color': '#667eea',
        'align': 'center',
        'valign': 'vcenter',
        'border': 1
    })
    
    warning_format = workbook.add_format({
        'bold': True,
        'font_color': '#78350f',
        'bg_color': '#fef3c7',
        'align': 'center',
        'text_wrap': True
    })
    
    # Largeurs colonnes
    worksheet_data.set_column('A:A', 25)
    worksheet_data.set_column('B:B', 25)
    worksheet_data.set_column('C:C', 18)
    worksheet_data.set_column('D:D', 28)
    worksheet_data.set_column('E:E', 28)
    worksheet_data.set_column('F:F', 20)
    worksheet_data.set_column('G:G', 20)
    worksheet_data.set_column('H:H', 30)
    
    # Message ligne 1
    worksheet_data.merge_range('A1:H1', 
        '📋 Remplissez vos données ci-dessous (le SEGMENT sera calculé à l\'import). Consultez le fichier Instructions (PDF) téléchargeable sur la page d\'accueil.',
        warning_format
    )
    worksheet_data.set_row(0, 30)
    
    # En-têtes ligne 2
    headers = [
        'NOM',
        'SECTEUR',
        'CA_HONORAIRES_HT',
        'OUTIL_COMPATIBLE_REFORME',
        'APPETENCE_INFORMATIQUE',
        'DIRIGEANT_PRENOM',
        'DIRIGEANT_NOM',
        'DIRIGEANT_EMAIL'
    ]
    
    for col_num, header in enumerate(headers):
        worksheet_data.write(1, col_num, header, header_format)
    
    # Données exemples
    examples = [
        ['SCI EXEMPLE', 'Immobilier', 8500, 'NON', 'TRES BON', 'Camille', 'DURAND', 'c.durand@exemple.fr'],
        ['SARL TEST', 'Hôtellerie', 4200, 'PARTIELLEMENT', 'BON', 'Marc', 'MARTIN', ''],
        ['EURL DEMO', 'Commerce de détail', 2100, 'OUI', 'MOYEN', '', '', '']
    ]
    
    for row_num, row_data in enumerate(examples, start=2):
        for col_num, cell_data in enumerate(row_data):
            worksheet_data.write(row_num, col_num, cell_data)
    
    # Validations
    worksheet_data.data_validation('D3:D1000', {
        'validate': 'list',
        'source': ['OUI', 'PARTIELLEMENT', 'NON'],
        'error_type': 'stop',
        'error_title': 'Valeur invalide',
        'error_message': 'Vous devez choisir : OUI, PARTIELLEMENT ou NON'
    })
    
    worksheet_data.data_validation('E3:E1000', {
        'validate': 'list',
        'source': ['TRES BON', 'BON', 'MOYEN', 'FAIBLE'],
        'error_type': 'stop',
        'error_title': 'Valeur invalide',
        'error_message': 'Vous devez choisir : TRES BON, BON, MOYEN ou FAIBLE'
    })
    
    worksheet_data.data_validation('C3:C1000', {
        'validate': 'decimal',
        'criteria': '>=',
        'value': 0,
        'error_type': 'warning',
        'error_title': 'Valeur suspecte',
        'error_message': 'Le CA doit être un nombre positif'
    })
    
    workbook.close()
    output.seek(0)
    
    return output.getvalue()


# ========================================
# FONCTION : CRÉATION FICHIER INSTRUCTIONS (PDF, 1 page A4)
# ========================================
def _creer_instructions_pdf_bytes():
    """Génère le guide d'utilisation au format PDF (une page A4)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return None
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=1.2*cm, bottomMargin=1.2*cm)
    styles = getSampleStyleSheet()
    
    c_titre = HexColor('#1e3a5f')
    c_section = HexColor('#2d5a87')
    c_sous = HexColor('#475569')
    c_texte = HexColor('#334155')
    c_footer = HexColor('#64748b')
    
    title_style = ParagraphStyle(
        name='CustomTitle', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=14, textColor=c_titre,
        alignment=1, spaceAfter=2, spaceBefore=0
    )
    subtitle_style = ParagraphStyle(
        name='Subtitle', parent=styles['Normal'],
        fontName='Helvetica', fontSize=9, textColor=c_sous, alignment=1, spaceAfter=10
    )
    heading_style = ParagraphStyle(
        name='CustomHeading', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=10, textColor=c_section,
        spaceBefore=8, spaceAfter=3
    )
    subheading_style = ParagraphStyle(
        name='SubHeading', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=9, textColor=c_section, spaceBefore=4, spaceAfter=2
    )
    body_style = ParagraphStyle(
        name='CustomBody', parent=styles['Normal'],
        fontName='Helvetica', fontSize=8, textColor=c_texte, leading=11, spaceAfter=2
    )
    body_indent = ParagraphStyle(
        name='BodyIndent', parent=body_style, leftIndent=18, spaceAfter=1
    )
    body_indent2 = ParagraphStyle(
        name='BodyIndent2', parent=body_style, leftIndent=28, spaceAfter=1
    )
    footer_style = ParagraphStyle(
        name='Footer', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=8, textColor=c_footer, spaceBefore=6, spaceAfter=2
    )
    footer_body = ParagraphStyle(
        name='FooterBody', parent=styles['Normal'],
        fontName='Helvetica', fontSize=8, textColor=c_footer, leading=10, spaceAfter=0
    )
    
    story = []
    story.append(Paragraph('GUIDE D\'UTILISATION \u2014 ANALYSE RFE 2026', title_style))
    story.append(Paragraph('Outil de Segmentation Client \u2014 Version 7.3', subtitle_style))
    
    story.append(Paragraph('1. COLONNES OBLIGATOIRES', heading_style))
    story.append(Paragraph('Ces données sont indispensables au fonctionnement de la matrice d\'analyse.', body_style))
    story.append(Paragraph('<b>IDENTIFICATION &amp; CHIFFRES</b>', subheading_style))
    story.append(Paragraph('<b>NOM</b> : Raison sociale du client.', body_indent))
    story.append(Paragraph('<b>SECTEUR</b> : Secteur d\'activité principal (texte libre, voir exemples en bas de page).', body_indent))
    story.append(Paragraph('<b>CA_HONORAIRES_HT</b> : Chiffre d\'affaires honoraires annuel en euros.', body_indent))
    story.append(Paragraph('Format : Numérique uniquement.', body_indent2))
    story.append(Paragraph('Décimales : Point . ou virgule , acceptés. Ne pas saisir de symbole &#8364; ni d\'espaces.', body_indent2))
    story.append(Paragraph('<b>DIAGNOSTIC (Listes déroulantes)</b>', subheading_style))
    story.append(Paragraph('Ces colonnes disposent d\'un menu déroulant (Validation de données). Merci de sélectionner l\'une des options proposées sans les modifier.', body_indent))
    story.append(Paragraph('<b>OUTIL_COMPATIBLE_REFORME</b>', body_indent))
    story.append(Paragraph('Objectif : Évaluer la conformité du logiciel de facturation actuel.', body_indent2))
    story.append(Paragraph('Choix : OUI | PARTIELLEMENT | NON', body_indent2))
    story.append(Paragraph('<b>APPETENCE_INFORMATIQUE</b>', body_indent))
    story.append(Paragraph('Objectif : Estimer le niveau d\'autonomie numérique pour le futur accompagnement.', body_indent2))
    story.append(Paragraph('Choix : TRES BON | BON | MOYEN | FAIBLE', body_indent2))
    
    story.append(Paragraph('2. COLONNES FACULTATIVES', heading_style))
    story.append(Paragraph('Ces champs permettent l\'automatisation des livrables.', body_style))
    story.append(Paragraph('<b>DIRIGEANT_PRENOM | DIRIGEANT_NOM | DIRIGEANT_EMAIL</b>', body_indent))
    story.append(Paragraph('Usage : Ces informations servent uniquement à personnaliser les modèles de courriers et emails de sensibilisation générés par l\'outil.', body_indent2))
    
    story.append(Paragraph('3. EXEMPLES DE SECTEURS', heading_style))
    story.append(Paragraph('Pour une meilleure lisibilité des statistiques, privilégiez les libellés suivants :', body_style))
    secteurs = 'Agriculture \u2022 Industrie et fabrication \u2022 Travaux de construction \u2022 Commerce de détail \u2022 Hôtellerie \u2022 Réparation de véhicules \u2022 Activités financières \u2022 Immobilier \u2022 Activités juridiques et comptables \u2022 Santé \u2022 Autres services aux personnes \u2022 Autres.'
    story.append(Paragraph(secteurs, body_indent))
    
    story.append(Paragraph('SÉCURITÉ &amp; CONFIDENTIALITÉ', footer_style))
    story.append(Paragraph('Traitement 100&#160;% local.', footer_body))
    story.append(Paragraph('L\'outil fonctionne en circuit fermé sur votre poste. Aucune donnée ne transite vers des serveurs externes.', footer_body))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def creer_instructions_pdf():
    """Retourne les octets du PDF d'instructions ou None si reportlab absent."""
    try:
        return _creer_instructions_pdf_bytes()
    except Exception:
        return None


# ========================================
# VÉRIFICATION DONNÉES CHARGÉES
# ========================================
if "df" not in st.session_state:
    st.session_state.df = None
if "kpis" not in st.session_state:
    st.session_state.kpis = None

# ========================================
# PAGE : ACCUEIL
# ========================================
if page == "🏠 Accueil":
    st.markdown(render_premium_header(
        title="🏆 Outil d'Analyse & Segmentation Client RFE",
        subtitle="Pilotez votre portefeuille clients avec intelligence - Réforme Facturation Électronique",
        badge="🔒 Traitement 100% local et confidentiel"
    ), unsafe_allow_html=True)
    
    st.markdown(section_divider("📋 COMMENCER"), unsafe_allow_html=True)
    
    col_left1, col_center1, col_right1 = st.columns([1, 2, 1])
    with col_center1:
        st.markdown("### 📥 Téléchargez le modèle et les instructions")
        st.markdown("""
        **Modèle Excel** :
        - ✅ Feuille "Données" : colonnes pré-configurées avec listes de choix
        - ✅ 3 lignes d'exemples pour comprendre le format
        - ✅ À remplir puis à importer ici
        
        **Fichier Instructions** (PDF) : guide pour remplir le modèle.
        """)
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                label="📥 Modèle Excel",
                data=creer_template_excel(),
                file_name="modele_clients_fae_2026.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="dl_template_accueil"
            )
        with col_dl2:
            instructions_pdf = creer_instructions_pdf()
            if instructions_pdf:
                st.download_button(
                    label="📄 Instructions (PDF)",
                    data=instructions_pdf,
                    file_name="Instruction modele Excel RFE.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key="dl_instructions_accueil"
                )
            else:
                st.caption("Instructions : installez reportlab pour générer le PDF.")
    
    st.markdown(section_divider("📤 IMPORTER VOS DONNÉES"), unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Sélectionner votre fichier d'import",
        type=["xlsx"],
        label_visibility="visible",
        key="accueil_upload"
    )
    
    if uploaded_file is not None:
        try:
            analyzer = FAEAnalyzer(coefficients, tarifs)
            df = analyzer.load_data(uploaded_file)
            
            if len(df) == 0:
                st.error("❌ Aucun client valide trouvé. Vérifiez que la colonne NOM est remplie.")
            else:
                df_scored = analyzer.calculate_scores(df)
                st.session_state.df = df_scored
                st.session_state.kpis = analyzer.calculate_advanced_kpis(df_scored)
                
                st.success(f"✅ {len(df_scored)} clients importés et analysés avec succès !")
                
                st.markdown(section_divider("📊 VUE D'ENSEMBLE"), unsafe_allow_html=True)
                
                col1, col2, col3, col4 = st.columns(4)
                
                kpis = st.session_state.kpis
                
                with col1:
                    st.markdown(metric_card_html(
                        icon="💰",
                        label="CA Total",
                        value=f"{kpis['ca_total']:,} €".replace(",", " "),
                        delta="+12%" if kpis['ca_total'] > 0 else None
                    ), unsafe_allow_html=True)
                
                with col2:
                    st.markdown(metric_card_html(
                        icon="👥",
                        label="Clients Actifs",
                        value=f"{kpis['total_clients']}",
                        delta=f"+{len(df_scored) // 20}" if len(df_scored) > 20 else None
                    ), unsafe_allow_html=True)
                
                with col3:
                    st.markdown(metric_card_html(
                        icon="⏰",
                        label="Jours avant réforme",
                        value=f"{kpis['jours_avant_reforme']} j",
                        delta=None
                    ), unsafe_allow_html=True)
                
                with col4:
                    st.markdown(metric_card_html(
                        icon="🎯",
                        label="Clients Prioritaires",
                        value=f"{kpis['tier1_count'] + kpis['tier2_count']}",
                        delta=None
                    ), unsafe_allow_html=True)
                
                st.markdown("---")
                st.info("👉 Utilisez le menu latéral pour accéder aux analyses détaillées")
                
        except Exception as e:
            st.error(f"❌ Erreur lors du chargement : {str(e)}")
            st.info("💡 Vérifiez que votre fichier respecte le format du modèle")
    
    else:
        st.markdown(section_divider("✨ FONCTIONNALITÉS"), unsafe_allow_html=True)
        
        col_feat1, col_feat2, col_feat3 = st.columns(3)
        
        with col_feat1:
            st.markdown("""
            ### 📊 Dashboard 360°
            - Vision complète de votre portefeuille
            - KPIs temps réel
            - Graphiques interactifs
            - Alertes automatiques
            """)
        
        with col_feat2:
            st.markdown("""
            ### 🎯 Priorisation IA
            - Scoring automatique clients
            - Segmentation intelligente
            - Actions recommandées
            - Potentiel CA calculé
            """)
        
        with col_feat3:
            st.markdown("""
            ### 📄 Livrables Pro
            - Rapports Word modifiables
            - Exports Excel enrichis
            - Modèles emails
            - Guides méthodologiques
            """)
# ========================================
# PAGE : DASHBOARD
# ========================================
elif page == "📊 Dashboard":
    st.markdown(render_premium_header(
        title="📊 Dashboard 360°",
        subtitle="Vue complète de votre portefeuille clients - Pilotage stratégique réforme FAE 2026"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        kpis = st.session_state.kpis
        
        # Section 1 : Urgence & Charge
        st.markdown(section_divider("⏰ URGENCE & CHARGE DE TRAVAIL"), unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("⏰ Jours avant réforme", f"{kpis['jours_avant_reforme']} jours", f"{kpis['mois_avant_reforme']} mois")
        
        with col2:
            st.metric("⏱️ Heures missions totales", f"{kpis['total_heures_mission']} h", "Toutes priorités")
        
        with col3:
            st.metric("👥 ETP nécessaires", f"{kpis['etp_necessaires']}", "Équivalent temps plein")
        
        with col4:
            st.metric("📞 Contacts/mois", f"{kpis['clients_a_traiter_par_mois']}", f"Sur {kpis['mois_avant_reforme']} mois")
        
        # Section 2 : État du portefeuille
        st.markdown(section_divider("📊 ÉTAT DU PORTEFEUILLE"), unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("👥 Clients totaux", f"{kpis['total_clients']}")
        
        with col2:
            pct_non_conforme = kpis['pct_flotte_non_conforme']
            st.metric("⚠️ Outils non conformes", f"{kpis['clients_outils_non_conformes']}", f"{pct_non_conforme}% du total")
        
        with col3:
            st.metric("🎯 Score maturité digitale", f"{kpis['score_maturite_digitale']}/100", "Moyenne portefeuille")
        
        with col4:
            st.metric("📅 Rythme requis", f"{kpis['clients_a_traiter_par_mois']} clients/mois", f"{int(kpis['clients_a_traiter_par_mois']/4)} par semaine")
        
        # Section 3 : Impact financier
        st.markdown(section_divider("💰 IMPACT FINANCIER"), unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("💼 CA cabinet actuel", f"{kpis['ca_total']:,} €".replace(",", " "))
        
        with col2:
            pct_risque = kpis['pct_ca_a_risque']
            st.metric("⚠️ CA à risque", f"{kpis['ca_total_a_risque']:,} €".replace(",", " "), f"{pct_risque}% du CA")
        
        with col3:
            st.metric("💎 Potentiel missions FAE", f"{kpis['ca_additionnel_max']:,} €".replace(",", " "), "Fourchette haute")
        
        with col4:
            st.metric("📈 Valeur moyenne/dossier", f"{kpis['valeur_fae_par_dossier']} €", "Potentiel unitaire")
        
        # Section 4 : Répartition priorités
        st.markdown(section_divider("🎯 RÉPARTITION PAR PRIORITÉ"), unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("🟢 PRIORITÉ 1", f"{kpis['tier1_count']} clients", "Audit complet")
        
        with col2:
            st.metric("🔵 PRIORITÉ 2", f"{kpis['tier2_count']} clients", "Formation")
        
        with col3:
            st.metric("🟠 PRIORITÉ 3", f"{kpis['tier3_count']} clients", "Information")
        
        with col4:
            st.metric("⚪ PRIORITÉ 4", f"{kpis['ignorer_count']} clients", "À surveiller")
        
        # Graphiques
        st.markdown(section_divider("📈 VISUALISATIONS"), unsafe_allow_html=True)
        
        col_left, col_right = st.columns(2)
        
        with col_left:
            st.markdown("#### 📊 Répartition par priorité")
            
            priorite_counts = df['PRIORITE'].value_counts()
            
            colors = {
                'PRIORITÉ 1 - Audit Complet': '#10b981',
                'PRIORITÉ 2 - Formation': '#3b82f6',
                'PRIORITÉ 3 - Information': '#f59e0b',
                'PRIORITÉ 4 - À Surveiller': '#9ca3af'
            }
            
            fig_pie = go.Figure(data=[go.Pie(
                labels=priorite_counts.index,
                values=priorite_counts.values,
                hole=0.4,
                marker=dict(colors=[colors.get(p, '#667eea') for p in priorite_counts.index]),
                textinfo='label+percent',
                textfont=dict(size=12)
            )])
            
            fig_pie.update_layout(
                showlegend=True,
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#ffffff', size=12),
                legend=dict(
                    font=dict(color='#ffffff', size=12),
                    bgcolor='rgba(0,0,0,0)',
                    bordercolor='rgba(255,255,255,0.2)',
                    borderwidth=1
                )
            )
            fig_pie.update_traces(textfont=dict(color='#ffffff', size=12))
            
            st.plotly_chart(fig_pie, use_container_width=True)
        
        with col_right:
            st.markdown("#### 💰 CA vs Score Opportunité")
            
            fig_scatter = px.scatter(
                df,
                x='CA_HONORAIRES_HT',
                y='SCORE_OPPORTUNITE',
                color='PRIORITE',
                size='CA_HONORAIRES_HT',
                hover_data=['NOM', 'SECTEUR'],
                color_discrete_map={
                    'PRIORITÉ 1 - Audit Complet': '#10b981',
                    'PRIORITÉ 2 - Formation': '#3b82f6',
                    'PRIORITÉ 3 - Information': '#f59e0b',
                    'PRIORITÉ 4 - À Surveiller': '#9ca3af'
                }
            )
            
            fig_scatter.update_layout(
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(15,23,42,0.8)',
                font=dict(color='#ffffff', size=12),
                legend=dict(
                    font=dict(color='#ffffff', size=12),
                    bgcolor='rgba(0,0,0,0)',
                    bordercolor='rgba(255,255,255,0.2)',
                    borderwidth=1
                ),
                xaxis=dict(
                    title="CA Honoraires (€)",
                    title_font=dict(color='#ffffff'),
                    tickfont=dict(color='#ffffff'),
                    gridcolor='rgba(255,255,255,0.15)',
                    zerolinecolor='rgba(255,255,255,0.2)'
                ),
                yaxis=dict(
                    title="Score Opportunité",
                    title_font=dict(color='#ffffff'),
                    tickfont=dict(color='#ffffff'),
                    gridcolor='rgba(255,255,255,0.15)',
                    zerolinecolor='rgba(255,255,255,0.2)'
                )
            )
            
            st.plotly_chart(fig_scatter, use_container_width=True)
        
        # Top 10 clients
        st.markdown(section_divider("🏆 TOP 10 CLIENTS À FORT POTENTIEL"), unsafe_allow_html=True)
        
        top10 = df.nlargest(10, 'SCORE_OPPORTUNITE')[['NOM', 'SECTEUR', 'CA_HONORAIRES_HT', 'SCORE_OPPORTUNITE', 'PRIORITE', 'ETOILES']]
        
        st.dataframe(
            top10,
            use_container_width=True,
            column_config={
                "NOM": st.column_config.TextColumn("Client", width="medium"),
                "SECTEUR": st.column_config.TextColumn("Secteur", width="medium"),
                "CA_HONORAIRES_HT": st.column_config.NumberColumn("CA (€)", format="%d €"),
                "SCORE_OPPORTUNITE": st.column_config.NumberColumn("Score", format="%.1f"),
                "PRIORITE": st.column_config.TextColumn("Priorité", width="large"),
                "ETOILES": st.column_config.TextColumn("⭐", width="small")
            },
            hide_index=True
        )

# ========================================
# PAGE : GUIDE MISSIONS
# ========================================
elif page == "🎯 Guide Missions":
    # Scroll en haut de page à l'arrivée sur Guide Missions
    st.markdown("""
        <div id="guide-missions-anchor"></div>
        <script>
            (function(){ var el = document.getElementById("guide-missions-anchor");
            if (el) el.scrollIntoView({ behavior: "instant", block: "start" }); })();
        </script>
    """, unsafe_allow_html=True)
    st.markdown(render_premium_header(
        title="🎯 Guide Missions FAE 2026",
        subtitle="Méthodologie d'accompagnement clients - 3 types de missions adaptées"
    ), unsafe_allow_html=True)
    
    st.markdown("""
        <div style="
            background: linear-gradient(135deg, rgba(102, 126, 234, 0.15) 0%, rgba(59, 130, 246, 0.15) 100%);
            border: 2px solid rgba(102, 126, 234, 0.5);
            border-radius: 16px;
            padding: 1.5rem;
            backdrop-filter: blur(10px);
        ">
            <h4 style="margin: 0 0 0.75rem 0; color: #60a5fa; font-weight: 700;">💡 Conseil cabinet</h4>
            <p style="margin: 0; font-size: 0.95rem; line-height: 1.5; color: #cbd5e1; font-weight: 400;">
                Commencez par analyser vos 20 plus gros clients. 
                Vous pouvez ensuite élargir à l'ensemble du portefeuille.
            </p>
        </div>
        """, unsafe_allow_html=True)
    
    # Mission 1
    with st.expander("🎯 MISSION 1 : Audit & Pilotage (Enjeu Élevé)", expanded=False):
        st.markdown("### Objectif")
        st.markdown("""
        Accompagner le client sur **tous les aspects critiques** de la facturation électronique : 
        diagnostic complet, mise en conformité outil, paramétrage, formation équipes et suivi post-démarrage.
        """)
        
        st.markdown("### Public cible")
        st.markdown("**PRIORITÉ 1** : Clients avec outil non conforme + CA significatif (>5000€)")
        
        st.markdown("### Durée & Format")
        st.markdown("- **8 heures** réparties sur **8 semaines**")
        st.markdown("- Format : 4 RDV de 2h (présentiel ou visio)")
        
        st.markdown("### Tarif")
        st.markdown("**1200€ à 1500€ HT** selon taille du dossier")
        
        st.markdown("### Déroulé détaillé")
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            st.markdown("**Phase 1**")
            st.markdown("Semaines 1-2")
        with col2:
            st.markdown("**Diagnostic & Préparation**")
            st.markdown("""
            - Audit de l'existant (outil actuel, processus, volumétrie factures)
            - Cartographie des flux (clients, fournisseurs, plateformes)
            - Identification des points de blocage
            - **Livrable :** Rapport d'audit (5 pages) + plan d'action
            """)
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            st.markdown("**Phase 2**")
            st.markdown("Semaines 3-5")
        with col2:
            st.markdown("**Mise en Conformité**")
            st.markdown("""
            - Choix/paramétrage outil conforme (accompagnement décision si changement)
            - Configuration : annuaires, mentions obligatoires, circuits de validation
            - Tests émission/réception factures
            - **Livrable :** Guide de paramétrage personnalisé
            """)
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            st.markdown("**Phase 3**")
            st.markdown("Semaine 6")
        with col2:
            st.markdown("**Formation Équipes**")
            st.markdown("""
            - Session formation 2h (comptable + dirigeant)
            - Cas pratiques : émettre/recevoir/corriger factures
            - Procédures anti-rejet (mentions manquantes, formats, etc.)
            - **Livrable :** Support formation + check-list quotidienne
            """)
        
        col1, col2 = st.columns([1, 3])
        
        with col1:
            st.markdown("**Phase 4**")
            st.markdown("Semaines 7-8")
        with col2:
            st.markdown("**Suivi & Ajustements**")
            st.markdown("""
            - Point hebdomadaire (15 min) : traitement anomalies
            - Vérification conformité premiers flux réels
            - Ajustements process si besoin
            - **Livrable :** Bilan final + recommandations pérennisation
            """)
        
        st.markdown("---")
        st.markdown("### 📧 Email type (Priorité 1)")
        
        email_p1 = """
Objet : [URGENT] Facturation Électronique 2026 - Audit de votre situation

Bonjour [Prénom],

La facturation électronique devient obligatoire le 1er septembre 2026. Selon notre analyse, votre dossier nécessite une attention particulière pour éviter tout blocage.

**Votre situation actuelle :**
- Outil de facturation : non conforme à la réforme
- Enjeu : continuité de votre activité (émission/réception factures)
- Échéance : 6 mois pour mettre en conformité

**Ce que nous vous proposons :**

✅ Audit complet de votre situation (outils, process, volumétrie)
✅ Accompagnement au choix/paramétrage d'une solution conforme
✅ Formation de vos équipes (comptable + vous-même)
✅ Suivi personnalisé post-démarrage (8 semaines)

📅 Je vous propose un RDV de 30 minutes cette semaine pour :
- Confirmer votre situation
- Identifier les ajustements nécessaires
- Vous présenter notre plan d'action simple

Êtes-vous disponible mardi 14h ou jeudi 10h ?

Cordialement,
[Signature]

P.S. : Cette mission est facturée entre 1200€ et 1500€ HT selon la complexité. Un investissement déductible qui sécurise votre activité.
        """
        
        st.text_area("", email_p1, height=400)
        st.download_button("📥 Télécharger l'email", email_p1, "email_priorite1.txt", use_container_width=True)

    # Mission 2
    with st.expander("🎓 MISSION 2 : Formation & Mise en Route (Enjeu Modéré)", expanded=False):
        st.markdown("### Objectif")
        st.markdown("""
        Rendre le client **autonome** sur la facturation électronique : comprendre les principes, 
        maîtriser l'outil (déjà conforme ou partiellement conforme), et éviter les erreurs courantes.
        """)
        
        st.markdown("### Public cible")
        st.markdown("**PRIORITÉ 2** : Clients avec outil partiellement conforme OU outil conforme mais faible appétence informatique")
        
        st.markdown("### Durée & Format")
        st.markdown("- **3 heures** sur **1 session** (ou 2× 1h30)")
        st.markdown("- Format : Atelier pratique (présentiel ou visio)")
        
        st.markdown("### Tarif")
        st.markdown("**600€ à 800€ HT**")
        
        st.markdown("### Contenu de la session")
        
        st.markdown("""
        **Partie 1 : Pré-diagnostic express (30 min)**
        - Revue rapide de l'outil actuel (est-il vraiment prêt ?)
        - Vérification des données de base (SIRET, adresses, TVA...)
        - Identification des 2-3 points d'attention prioritaires
        
        **Partie 2 : Formation action (1h30)**
        - Principes de la facturation électronique (ce qui change vraiment)
        - Cycle de vie d'une facture : émission, transmission, réception, archivage
        - Démonstration live : émettre une facture conforme
        - Cas pratique : le client émet sa première facture sous supervision
        
        **Partie 3 : Sécurisation & Bonnes pratiques (1h)**
        - Les 5 erreurs qui font rejeter une facture (et comment les éviter)
        - Checklist mensuelle de contrôle (à faire soi-même)
        - Que faire en cas de problème ? (support, ressources)
        - Q&R personnalisées
        """)
        
        st.markdown("### Livrables")
        st.markdown("""
        - Support de formation (PDF 15 pages)
        - Checklist anti-rejet (1 page A4)
        - Accès documentation vidéo (3 tutos 5 min)
        """)
        
        st.markdown("---")
        st.markdown("### 📧 Email type (Priorité 2)")
        
        email_p2 = """
Objet : Facturation Électronique 2026 - Formation pour votre équipe

Bonjour [Prénom],

Bonne nouvelle : votre outil de facturation est sur la bonne trajectoire pour la réforme de septembre 2026. 

**Où en êtes-vous ?**
Votre solution est conforme (ou nécessite juste une mise à jour), mais l'enjeu est maintenant de **rendre vos équipes autonomes** pour éviter les rejets de factures et les blocages administratifs.

**Ce que nous vous proposons :**

🎓 Session de formation pratique (3 heures)
- Comprendre les nouveaux principes (sans jargon technique)
- Manipuler votre outil en conditions réelles
- Éviter les 5 erreurs classiques qui bloquent tout

📋 Livrables inclus :
- Support de formation complet
- Checklist anti-rejet (à afficher dans le bureau)
- Vidéos tutos (3× 5 min pour se remettre à niveau si besoin)

📅 Plusieurs créneaux disponibles en mars :
- Mardi 12/03 : 14h-17h
- Jeudi 14/03 : 9h-12h
- Vendredi 15/03 : 14h-17h

Tarif : 600€ HT (déductible fiscalement)

Répondez à cet email pour réserver votre créneau (places limitées à 8 sessions/mois).

Cordialement,
[Signature]
        """
        
        st.text_area("", email_p2, height=350)
        st.download_button("📥 Télécharger l'email", email_p2, "email_priorite2.txt", use_container_width=True)

    # Mission 3
    with st.expander("📢 MISSION 3 : Information & Sensibilisation (Enjeu Faible)", expanded=False):
        st.markdown("### Objectif")
        st.markdown("""
        Informer le client sur les **essentiels** de la réforme, vérifier qu'il est sur les bons rails, 
        et le rassurer (pas de panique, c'est gérable).
        """)
        
        st.markdown("### Public cible")
        st.markdown("**PRIORITÉ 3** : Clients avec outil déjà conforme + bonne appétence informatique")
        
        st.markdown("### Durée & Format")
        st.markdown("- **30 à 60 minutes** (appel téléphonique ou visio)")
        st.markdown("- Format : Point de situation rapide")
        
        st.markdown("### Tarif")
        st.markdown("**150€ à 300€ HT** OU **inclus dans vos honoraires récurrents** (positionnement conseil)")
        
        st.markdown("### Contenu de l'échange")
        
        st.markdown("""
        1. **Confirmer la conformité outil** (5 min)
           - "Votre solution est bien compatible, voici pourquoi..."
        
        2. **Les 3 points de vigilance** (15 min)
           - Vérifier que vos coordonnées sont à jour (SIRET, adresse, email de facturation)
           - Suivre les mises à jour de votre outil (ne pas ignorer les notifications)
           - Tester l'émission/réception d'une facture avant le 1er septembre
        
        3. **Message rassurant** (10 min)
           - Vous êtes déjà bien positionné
           - Pas de changement majeur dans vos habitudes
           - On reste disponible si problème (hotline cabinet)
        
        4. **Mini-checklist envoyée par email** (post-appel)
           - 5 points à cocher d'ici septembre
        """)
        
        st.markdown("### Livrable")
        st.markdown("- Email récapitulatif avec mini-checklist (1 page)")
        
        st.markdown("---")
        st.markdown("### 📧 Email type (Priorité 3)")
        
        email_p3 = """
Objet : Facturation Électronique 2026 - Vous êtes prêt (presque !)

Bonjour [Prénom],

La réforme de la facturation électronique arrive en septembre 2026. Bonne nouvelle : selon notre analyse, vous êtes déjà sur les bons rails.

**Votre situation :**
✅ Outil de facturation conforme
✅ Bonne maîtrise des outils numériques
✅ Process en place

**Ce qu'il reste à faire (simple) :**

1. Vérifier vos coordonnées (SIRET, adresse, email facturation)
2. Suivre les mises à jour de votre logiciel (automne 2026)
3. Tester l'envoi d'une facture électronique (nous vous accompagnons si besoin)

📞 Je vous propose un point téléphonique de 30 minutes pour :
- Confirmer que tout est OK
- Répondre à vos questions éventuelles
- Vous donner la checklist finale

Disponible la semaine du [date] : mardi 10h, mercredi 14h ou jeudi 9h ?

Pas d'urgence, mais autant valider maintenant pour être serein en septembre.

Cordialement,
[Signature]

P.S. : Ce point est inclus dans nos honoraires habituels (logique conseil).
        """
        
        st.text_area("", email_p3, height=350)
        st.download_button("📥 Télécharger l'email", email_p3, "email_priorite3.txt", use_container_width=True)

# ========================================
# PAGE : MATRICE 2.0
# ========================================
elif page == "🔍 Matrice 2.0":
    st.markdown(render_premium_header(
        title="🔍 Matrice 2.0 - Segmentation Avancée",
        subtitle="Analyse multi-critères et filtres dynamiques"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        
        st.markdown(section_divider("🎛️ FILTRES DYNAMIQUES"), unsafe_allow_html=True)
        
        # Segment CA en premier, seul sur toute la largeur (textes longs / montants visibles)
        segment_filter = st.multiselect(
            "Segment CA",
            options=df['SEGMENT'].unique(),
            default=df['SEGMENT'].unique(),
            key="matrice_segment"
        )
        
        # Appétence et Conformité Outil sur une même ligne en dessous
        col_app, col_outil = st.columns(2)
        with col_app:
            appetence_filter = st.multiselect(
                "Appétence Informatique",
                options=df['APPETENCE_INFORMATIQUE'].unique(),
                default=df['APPETENCE_INFORMATIQUE'].unique(),
                key="matrice_appetence"
            )
        with col_outil:
            outil_filter = st.multiselect(
                "Conformité Outil",
                options=df['OUTIL_COMPATIBLE_REFORME'].unique(),
                default=df['OUTIL_COMPATIBLE_REFORME'].unique(),
                key="matrice_outil"
            )
        
        # Application filtres
        df_filtered = df[
            (df['APPETENCE_INFORMATIQUE'].isin(appetence_filter)) &
            (df['SEGMENT'].isin(segment_filter)) &
            (df['OUTIL_COMPATIBLE_REFORME'].isin(outil_filter))
        ]
        
        st.info(f"📊 {len(df_filtered)} clients correspondent aux filtres sélectionnés (sur {len(df)} total)")
        
        # Heatmaps
        st.markdown(section_divider("🔥 HEATMAPS CROISÉES"), unsafe_allow_html=True)
        
        col_heat1, col_heat2 = st.columns(2)
        
        with col_heat1:
            st.markdown("#### Appétence × Outil")
            
            pivot_app_outil = pd.crosstab(
                df_filtered['APPETENCE_INFORMATIQUE'],
                df_filtered['OUTIL_COMPATIBLE_REFORME']
            )
            
            # Texte = nombre de clients par cellule (selon filtres)
            text_app_outil = pivot_app_outil.values.astype(int).astype(str)
            
            fig_heat1 = go.Figure(data=go.Heatmap(
                z=pivot_app_outil.values,
                x=pivot_app_outil.columns.tolist(),
                y=pivot_app_outil.index.tolist(),
                text=text_app_outil,
                texttemplate="%{text}",
                textfont=dict(size=14, color="#ffffff"),
                colorscale='RdYlGn',
                colorbar=dict(title="Nb Clients"),
                hoverongaps=False
            ))
            
            fig_heat1.update_layout(
                height=350,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#ffffff'),
                xaxis=dict(title="Conformité Outil", title_font=dict(color='#ffffff'), tickfont=dict(color='#ffffff')),
                yaxis=dict(title="Appétence", title_font=dict(color='#ffffff'), tickfont=dict(color='#ffffff'))
            )
            
            st.plotly_chart(fig_heat1, use_container_width=True)
        
        with col_heat2:
            st.markdown("#### Segment CA × Outil")
            
            pivot_seg_outil = pd.crosstab(
                df_filtered['SEGMENT'],
                df_filtered['OUTIL_COMPATIBLE_REFORME']
            )
            
            text_seg_outil = pivot_seg_outil.values.astype(int).astype(str)
            
            fig_heat2 = go.Figure(data=go.Heatmap(
                z=pivot_seg_outil.values,
                x=pivot_seg_outil.columns.tolist(),
                y=pivot_seg_outil.index.tolist(),
                text=text_seg_outil,
                texttemplate="%{text}",
                textfont=dict(size=14, color="#1a1a1a", family="Arial Black, sans-serif"),
                colorscale='Blues',
                colorbar=dict(title="Nb Clients"),
                hoverongaps=False,
                hoverlabel=dict(bgcolor='#1e293b', font=dict(color='#ffffff'))
            ))
            
            fig_heat2.update_layout(
                height=350,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#ffffff'),
                xaxis=dict(title="Conformité Outil", title_font=dict(color='#ffffff'), tickfont=dict(color='#ffffff')),
                yaxis=dict(title="Segment CA", title_font=dict(color='#ffffff'), tickfont=dict(color='#ffffff'))
            )
            
            st.plotly_chart(fig_heat2, use_container_width=True)
        
        # Table détaillée
        st.markdown(section_divider("📋 TABLE DÉTAILLÉE"), unsafe_allow_html=True)
        
        df_display = df_filtered.sort_values('SCORE_OPPORTUNITE', ascending=False)[
            ['NOM', 'SECTEUR', 'CA_HONORAIRES_HT', 'OUTIL_COMPATIBLE_REFORME', 
             'APPETENCE_INFORMATIQUE', 'SCORE_OPPORTUNITE', 'PRIORITE', 'ETOILES']
        ]
        
        st.dataframe(
            df_display,
            use_container_width=True,
            column_config={
                "NOM": st.column_config.TextColumn("Client", width="medium"),
                "SECTEUR": st.column_config.TextColumn("Secteur", width="medium"),
                "CA_HONORAIRES_HT": st.column_config.NumberColumn("CA (€)", format="%d €"),
                "OUTIL_COMPATIBLE_REFORME": st.column_config.TextColumn("Outil", width="small"),
                "APPETENCE_INFORMATIQUE": st.column_config.TextColumn("Appétence", width="small"),
                "SCORE_OPPORTUNITE": st.column_config.NumberColumn("Score", format="%.1f"),
                "PRIORITE": st.column_config.TextColumn("Priorité", width="large"),
                "ETOILES": st.column_config.TextColumn("⭐", width="small")
            },
            hide_index=True
        )
        
        # Export sélection
        st.markdown("---")
        
        col_exp1, col_exp2, col_exp3 = st.columns([1, 1, 2])
        
        with col_exp1:
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df_display.to_excel(writer, sheet_name='Sélection filtrée', index=False)
            
            st.download_button(
                label="📥 Exporter la sélection (Excel)",
                data=output.getvalue(),
                file_name=f"selection_filtree_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
# ========================================
# PAGE : SIMULATEUR CA
# ========================================
elif page == "💰 Simulateur CA":
    st.markdown(render_premium_header(
        title="💰 Simulateur CA Missions FAE",
        subtitle="Projections financières et scénarios de conversion"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        kpis = st.session_state.kpis
        
        st.markdown("""
        <div class="mode-emploi-box" style="
            background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%);
            border: 2px solid #f59e0b;
            border-radius: 16px;
            padding: 1.5rem;
            color: #1a1a1a;
            margin-bottom: 2rem;
        ">
            <h4 style="margin: 0 0 0.75rem 0; color: #1a1a1a;">💡 Mode d'emploi</h4>
            <p style="margin: 0; font-size: 0.95rem; line-height: 1.5; color: #1a1a1a;">
                Ajustez les curseurs ci-dessous pour simuler différents scénarios de conversion.
                Le CA additionnel se calcule en temps réel selon vos hypothèses.
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(section_divider("🎯 PARAMÈTRES PAR SEGMENT"), unsafe_allow_html=True)
        
        # Priorité 1
        st.markdown("### 🟢 PRIORITÉ 1 - Audit Complet")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            p1_active = st.checkbox("Activer P1", value=True, key="p1_active")
        
        with col2:
            p1_taux = st.slider("Taux de conversion", 0, 100, 70, 5, key="p1_taux", disabled=not p1_active)
        
        with col3:
            p1_prix = st.number_input("Prix unitaire (€)", 1000, 2000, 1350, 50, key="p1_prix", disabled=not p1_active)
        
        # Priorité 2
        st.markdown("### 🔵 PRIORITÉ 2 - Formation")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            p2_active = st.checkbox("Activer P2", value=True, key="p2_active")
        
        with col2:
            p2_taux = st.slider("Taux de conversion", 0, 100, 60, 5, key="p2_taux", disabled=not p2_active)
        
        with col3:
            p2_prix = st.number_input("Prix unitaire (€)", 500, 1000, 700, 50, key="p2_prix", disabled=not p2_active)
        
        # Priorité 3
        st.markdown("### 🟠 PRIORITÉ 3 - Information")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            p3_active = st.checkbox("Activer P3", value=True, key="p3_active")
        
        with col2:
            p3_taux = st.slider("Taux de conversion", 0, 100, 40, 5, key="p3_taux", disabled=not p3_active)
        
        with col3:
            p3_prix = st.number_input("Prix unitaire (€)", 100, 500, 225, 25, key="p3_prix", disabled=not p3_active)
        
        # Calculs
        st.markdown(section_divider("💰 RÉSULTATS"), unsafe_allow_html=True)
        
        p1_count = kpis['tier1_count']
        p2_count = kpis['tier2_count']
        p3_count = kpis['tier3_count']
        
        ca_p1 = (p1_count * (p1_taux / 100) * p1_prix) if p1_active else 0
        ca_p2 = (p2_count * (p2_taux / 100) * p2_prix) if p2_active else 0
        ca_p3 = (p3_count * (p3_taux / 100) * p3_prix) if p3_active else 0
        
        ca_total_simule = ca_p1 + ca_p2 + ca_p3
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("🟢 CA P1", f"{int(ca_p1):,} €".replace(",", " "))
        
        with col2:
            st.metric("🔵 CA P2", f"{int(ca_p2):,} €".replace(",", " "))
        
        with col3:
            st.metric("🟠 CA P3", f"{int(ca_p3):,} €".replace(",", " "))
        
        with col4:
            pct_ca = (ca_total_simule / kpis['ca_total'] * 100) if kpis['ca_total'] > 0 else 0
            st.metric("💎 CA TOTAL", f"{int(ca_total_simule):,} €".replace(",", " "), f"{pct_ca:.1f}% du CA cabinet")
        
        # Détail conversions
        st.markdown(section_divider("📋 DÉTAIL CONVERSIONS"), unsafe_allow_html=True)
        
        data_simulation = {
            "Priorité": ["P1 - Audit", "P2 - Formation", "P3 - Information"],
            "Clients éligibles": [p1_count, p2_count, p3_count],
            "Taux conversion": [f"{p1_taux}%" if p1_active else "0%", 
                                f"{p2_taux}%" if p2_active else "0%", 
                                f"{p3_taux}%" if p3_active else "0%"],
            "Clients convertis": [int(p1_count * p1_taux / 100) if p1_active else 0,
                                   int(p2_count * p2_taux / 100) if p2_active else 0,
                                   int(p3_count * p3_taux / 100) if p3_active else 0],
            "Prix unitaire": [f"{p1_prix}€" if p1_active else "0€",
                              f"{p2_prix}€" if p2_active else "0€",
                              f"{p3_prix}€" if p3_active else "0€"],
            "CA généré": [f"{int(ca_p1):,}€".replace(",", " "),
                          f"{int(ca_p2):,}€".replace(",", " "),
                          f"{int(ca_p3):,}€".replace(",", " ")]
        }
        
        df_simulation = pd.DataFrame(data_simulation)
        
        st.dataframe(df_simulation, use_container_width=True, hide_index=True)

# ========================================
# PAGE : PLAN D'ACTION
# ========================================
elif page == "📅 Plan d'Action":
    st.markdown(render_premium_header(
        title="📅 Plan d'Action - Roadmap 6 Mois",
        subtitle="Organisation commerciale et planning déploiement"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        kpis = st.session_state.kpis
        
        # Rythme recommandé
        st.markdown(section_divider("⏱️ RYTHME RECOMMANDÉ"), unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("📞 Contacts/mois", f"{kpis['clients_a_traiter_par_mois']}", "Moyenne à maintenir")
        
        with col2:
            contacts_semaine = kpis['clients_a_traiter_par_mois'] / 4
            st.metric("📅 Contacts/semaine", f"{int(contacts_semaine)}", "Soit ~2-3 par jour")
        
        with col3:
            heures_mois = kpis['total_heures_mission'] / kpis['mois_avant_reforme']
            st.metric("⏰ Heures/mois", f"{int(heures_mois)} h", "Charge missions")
        
        # Planning Gantt
        st.markdown(section_divider("📊 PLANNING GANTT 6 MOIS"), unsafe_allow_html=True)
        
        today = datetime.now()
        
        # Ordre affiché : P1 en premier (haut), puis P2, P3, Phase 4 (Suivi) en bas
        phases = [
            dict(Phase="Phase 1 - Sécurisation P1", Start=today, Finish=today + timedelta(days=45), Resource="Priorité 1"),
            dict(Phase="Phase 2 - Montée P2", Start=today + timedelta(days=30), Finish=today + timedelta(days=90), Resource="Priorité 2"),
            dict(Phase="Phase 3 - Information P3", Start=today + timedelta(days=60), Finish=today + timedelta(days=120), Resource="Priorité 3"),
            dict(Phase="Phase 4 - Relances", Start=today + timedelta(days=90), Finish=today + timedelta(days=150), Resource="Suivi")
        ]
        
        df_gantt = pd.DataFrame(phases)
        
        color_map = {
            "Priorité 1": "#10b981",
            "Priorité 2": "#3b82f6",
            "Priorité 3": "#f59e0b",
            "Suivi": "#9ca3af"
        }
        
        fig_gantt = px.timeline(
            df_gantt,
            x_start="Start",
            x_end="Finish",
            y="Phase",
            color="Resource",
            color_discrete_map=color_map
        )
        
        # Premier élément = en bas du graphique, dernier = en haut → P1 en haut
        order_phase = ["Phase 4 - Relances", "Phase 3 - Information P3", "Phase 2 - Montée P2", "Phase 1 - Sécurisation P1"]
        fig_gantt.update_layout(
            height=400,
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(15,23,42,0.8)',
            font=dict(color='#ffffff', size=12),
            legend=dict(
                font=dict(color='#ffffff', size=12),
                bgcolor='rgba(0,0,0,0)',
                bordercolor='rgba(255,255,255,0.2)',
                borderwidth=1
            ),
            xaxis=dict(
                title="",
                title_font=dict(color='#ffffff'),
                tickfont=dict(color='#ffffff'),
                gridcolor='rgba(255,255,255,0.15)',
                zerolinecolor='rgba(255,255,255,0.2)'
            ),
            yaxis=dict(
                title="",
                title_font=dict(color='#ffffff'),
                tickfont=dict(color='#ffffff'),
                categoryorder='array',
                categoryarray=order_phase
            )
        )
        
        st.plotly_chart(fig_gantt, use_container_width=True)
        
        # Playbook hebdomadaire
        st.markdown(section_divider("📋 PLAYBOOK HEBDOMADAIRE TYPE"), unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("**Lundi**")
            st.markdown("🎯 Planification")
        with col2:
            st.markdown("""
            - Sélectionner 8-10 clients à contacter cette semaine
            - Préparer les emails/appels (modèles + personnalisation)
            - Bloquer 2h dans l'agenda pour envois groupés
            """)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("**Mardi-Jeudi**")
            st.markdown("📞 Exécution")
        with col2:
            st.markdown("""
            - Envoi emails (mardi 9h)
            - Relances téléphoniques J+2 (jeudi 14h)
            - RDV confirmés → planifier missions
            """)
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("**Vendredi**")
            st.markdown("📊 Reporting")
        with col2:
            st.markdown("""
            - Mettre à jour CRM (statuts clients)
            - Comptabiliser taux de conversion
            - Préparer liste semaine suivante
            """)
        
        # Extraction cible semaine
        st.markdown(section_divider("🎯 EXTRACTION CIBLE SEMAINE"), unsafe_allow_html=True)
        
        nb_clients_semaine = st.number_input(
            "Nombre de clients à contacter cette semaine",
            min_value=1,
            max_value=50,
            value=int(contacts_semaine) if contacts_semaine > 0 else 10
        )
        
        priorite_cible = st.selectbox(
            "Priorité ciblée",
            ["PRIORITÉ 1 - Audit Complet", "PRIORITÉ 2 - Formation", "PRIORITÉ 3 - Information", "Toutes priorités"]
        )
        
        if priorite_cible == "Toutes priorités":
            df_cible = df.nlargest(nb_clients_semaine, 'SCORE_OPPORTUNITE')
        else:
            df_cible = df[df['PRIORITE'] == priorite_cible].nlargest(nb_clients_semaine, 'SCORE_OPPORTUNITE')
        
        st.markdown(f"**👥 {len(df_cible)} clients sélectionnés**")
        
        df_cible_display = df_cible[['NOM', 'SECTEUR', 'CA_HONORAIRES_HT', 'PRIORITE', 'SCORE_OPPORTUNITE', 'DIRIGEANT_EMAIL']]
        
        st.dataframe(df_cible_display, use_container_width=True, hide_index=True)
        
        # Export cible
        output_cible = io.BytesIO()
        with pd.ExcelWriter(output_cible, engine='xlsxwriter') as writer:
            df_cible_display.to_excel(writer, sheet_name='Cibles semaine', index=False)
        
        st.download_button(
            label=f"📥 Télécharger la liste ({len(df_cible)} clients)",
            data=output_cible.getvalue(),
            file_name=f"cibles_semaine_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

# ========================================
# PAGE : ANALYSE RISQUES
# ========================================
elif page == "⚠️ Analyse Risques":
    st.markdown(render_premium_header(
        title="⚠️ Analyse Risques - Clients à Sécuriser",
        subtitle="Identification et priorisation des dossiers critiques"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        
        # Calcul score risque
        def calculate_risk_score(row):
            score = 0
            
            # CA élevé = risque élevé
            if row['CA_HONORAIRES_HT'] >= 10000:
                score += 40
            elif row['CA_HONORAIRES_HT'] >= 5000:
                score += 30
            elif row['CA_HONORAIRES_HT'] >= 2500:
                score += 20
            elif row['CA_HONORAIRES_HT'] >= 1000:
                score += 10
            
            # Outil non conforme = risque élevé
            if row['OUTIL_COMPATIBLE_REFORME'] == 'NON':
                score += 40
            elif row['OUTIL_COMPATIBLE_REFORME'] == 'PARTIELLEMENT':
                score += 20
            
            # Faible appétence = risque modéré
            if row['APPETENCE_INFORMATIQUE'] == 'FAIBLE':
                score += 20
            elif row['APPETENCE_INFORMATIQUE'] == 'MOYEN':
                score += 10
            
            return score
        
        df['SCORE_RISQUE'] = df.apply(calculate_risk_score, axis=1)
        
        def get_risk_level(score):
            if score >= 75:
                return "CRITIQUE"
            elif score >= 55:
                return "ÉLEVÉ"
            elif score >= 35:
                return "MOYEN"
            else:
                return "FAIBLE"
        
        df['NIVEAU_RISQUE'] = df['SCORE_RISQUE'].apply(get_risk_level)
        
        # Métriques
        st.markdown(section_divider("⚠️ VUE D'ENSEMBLE RISQUES"), unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        
        risque_critique = len(df[df['NIVEAU_RISQUE'] == 'CRITIQUE'])
        risque_eleve = len(df[df['NIVEAU_RISQUE'] == 'ÉLEVÉ'])
        risque_moyen = len(df[df['NIVEAU_RISQUE'] == 'MOYEN'])
        risque_faible = len(df[df['NIVEAU_RISQUE'] == 'FAIBLE'])
        
        with col1:
            st.metric("🔴 CRITIQUE", risque_critique, "Action immédiate")
        
        with col2:
            st.metric("🟠 ÉLEVÉ", risque_eleve, "Sous 2 semaines")
        
        with col3:
            st.metric("🟡 MOYEN", risque_moyen, "Surveillance")
        
        with col4:
            st.metric("🟢 FAIBLE", risque_faible, "Situation OK")
        
        # Distribution
        st.markdown(section_divider("📊 DISTRIBUTION SCORES RISQUE"), unsafe_allow_html=True)
        
        fig_hist = px.histogram(
            df,
            x='SCORE_RISQUE',
            color='NIVEAU_RISQUE',
            nbins=20,
            color_discrete_map={
                'CRITIQUE': '#dc2626',
                'ÉLEVÉ': '#ea580c',
                'MOYEN': '#f59e0b',
                'FAIBLE': '#10b981'
            }
        )
        
        fig_hist.update_layout(
            height=400,
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(15,23,42,0.8)',
            font=dict(color='#f1f5f9'),
            xaxis=dict(title="Score Risque", gridcolor='rgba(255,255,255,0.1)'),
            yaxis=dict(title="Nombre de clients", gridcolor='rgba(255,255,255,0.1)')
        )
        
        st.plotly_chart(fig_hist, use_container_width=True)
        
        # Top clients à risque
        st.markdown(section_divider("🚨 TOP CLIENTS À SÉCURISER"), unsafe_allow_html=True)
        
        nb_top_risque = st.slider("Nombre de clients à afficher", 5, 50, 20)
        
        df_risque = df.nlargest(nb_top_risque, 'SCORE_RISQUE')[
            ['NOM', 'SECTEUR', 'CA_HONORAIRES_HT', 'OUTIL_COMPATIBLE_REFORME', 
             'APPETENCE_INFORMATIQUE', 'SCORE_RISQUE', 'NIVEAU_RISQUE']
        ]
        
        st.dataframe(
            df_risque,
            use_container_width=True,
            column_config={
                "NOM": st.column_config.TextColumn("Client", width="medium"),
                "SECTEUR": st.column_config.TextColumn("Secteur", width="medium"),
                "CA_HONORAIRES_HT": st.column_config.NumberColumn("CA (€)", format="%d €"),
                "OUTIL_COMPATIBLE_REFORME": st.column_config.TextColumn("Outil", width="small"),
                "APPETENCE_INFORMATIQUE": st.column_config.TextColumn("Appétence", width="small"),
                "SCORE_RISQUE": st.column_config.NumberColumn("Score", format="%d"),
                "NIVEAU_RISQUE": st.column_config.TextColumn("Niveau", width="small")
            },
            hide_index=True
        )
        
        # Export
        output_risque = io.BytesIO()
        with pd.ExcelWriter(output_risque, engine='xlsxwriter') as writer:
            df_risque.to_excel(writer, sheet_name='Clients à risque', index=False)
        
        st.download_button(
            label=f"📥 Exporter les {nb_top_risque} clients à risque",
            data=output_risque.getvalue(),
            file_name=f"clients_risque_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

# ========================================
# PAGE : BIBLIOTHÈQUE
# ========================================
elif page == "📚 Bibliothèque":
    st.markdown(render_premium_header(
        title="📚 Bibliothèque de Contenus",
        subtitle="Modèles d'emails prêts à l'emploi"
    ), unsafe_allow_html=True)
    
    st.markdown(section_divider("📧 TEMPLATES EMAILS"), unsafe_allow_html=True)
    
    email_choice = st.selectbox(
        "Choisissez un modèle",
        [
            "Email 1 - Audit urgence (Priorité 1)",
            "Email 2 - Formation structurée (Priorité 2)",
            "Email 3 - Information simple (Priorité 3)",
            "Email 4 - Relance J+3 (sans réponse)",
            "Email 5 - Relance J+10 (dernière relance)"
        ]
    )
    
    emails_library = {
        "Email 1 - Audit urgence (Priorité 1)": """Objet : Facturation Électronique 2026 - Audit de votre situation

Bonjour [Prénom],

La facturation électronique devient obligatoire le 1er septembre 2026. Après analyse de votre dossier, votre situation nécessite une attention particulière pour assurer la continuité de votre activité.

Votre situation actuelle :

    Outil de facturation : non conforme à la réforme
    Enjeu : continuité d'émission/réception de vos factures
    Échéance : 6 mois pour la mise en conformité

Notre accompagnement :
✅ Audit complet (outils, process, volumétrie)

✅ Accompagnement au choix d'une solution adaptée

✅ Formation de vos équipes

✅ Suivi post-démarrage (8 semaines)

📅 Je vous propose un point de 30 minutes cette semaine ou la semaine prochaine pour :

    Confirmer votre situation
    Identifier les ajustements nécessaires
    Vous présenter un plan d'action

Cordialement,

[Signature]""",
        
        "Email 2 - Formation structurée (Priorité 2)": """Objet : Facturation Électronique 2026 - Formation de vos équipes

Bonjour [Prénom],

Bonne nouvelle : votre outil de facturation est compatible avec la réforme de septembre 2026.

L'enjeu maintenant : rendre vos équipes autonomes pour éviter les rejets de factures et les blocages administratifs.

Notre proposition :
🎓 Session de formation pratique (3 heures)

    Comprendre les nouveaux principes (sans jargon)
    Manipuler votre outil en conditions réelles
    Éviter les erreurs classiques

📋 Livrable inclus :

    Support de formation complet

📅 Je vous propose un point de 30 minutes cette semaine ou la semaine prochaine pour :

    Valider vos besoins
    Planifier la session de formation

Répondez à cet email pour convenir d'un échange.
Cordialement,

[Signature]""",
        
        "Email 3 - Information simple (Priorité 3)": """Objet : Facturation Électronique 2026 - Validation de votre conformité

Bonjour [Prénom],

Selon notre analyse, vous êtes bien positionné pour la réforme de septembre 2026.

Votre situation :
✅ Outil de facturation conforme

✅ Bonne maîtrise des outils numériques

✅ Process en place

Points de vigilance (simples) :

    Vérifier vos coordonnées (SIRET, adresse, email facturation)
    Suivre les mises à jour logiciel (automne 2026)
    Tester l'envoi d'une première facture électronique

📞 Je vous propose un point téléphonique de 30 minutes cette semaine ou la semaine prochaine pour :

    Confirmer que tout est en ordre
    Répondre à vos questions
    Vous remettre la checklist finale

Pas d'urgence, mais autant valider maintenant pour être serein en septembre.
Cordialement,

[Signature]

Ce point est inclus dans nos honoraires habituels.""",
        
        "Email 4 - Relance J+3 (sans réponse)": """Objet : Re: Facturation Électronique 2026

Bonjour [Prénom],

Je reviens vers vous suite à mon email de [jour].

La date du 1er septembre 2026 approche et certains dossiers nécessitent un délai de mise en conformité.

Rappel :

    Votre outil : [conforme/non conforme]
    Action recommandée : [audit/formation/validation]
    Délai estimé : [X semaines]

Je vous propose un point de 30 minutes cette semaine ou la semaine prochaine pour faire le point.
Si ce n'est pas le bon moment, dites-le-moi : je vous recontacterai dans 2 mois.
Cordialement,

[Signature]""",
        
        "Email 5 - Relance J+10 (dernière relance)": """Objet : [Dernière relance] Facturation Électronique 2026

Bonjour [Prénom],

C'est ma dernière relance sur ce sujet.

Mon rôle est de vous alerter sur une obligation réglementaire :
⚠️ [X] mois avant le 1er septembre 2026

⚠️ Votre outil actuel : non conforme (selon nos informations)

⚠️ Risque : blocage de facturation = impact direct sur votre activité

Si vous avez déjà traité le sujet avec un autre prestataire, merci de me le signaler afin que je clôture votre dossier.
Dans le cas contraire, je reste à votre disposition pour un échange cette semaine.
📞 [Téléphone]

📧 Réponse directe à cet email
Cordialement,

[Signature]"""
    }
    
    selected_email = emails_library[email_choice]
    
    st.text_area("Contenu de l'email", selected_email, height=450)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.download_button(
            "📥 Télécharger (.txt)",
            selected_email,
            file_name=f"{email_choice.replace(' ', '_').lower()}.txt",
            use_container_width=True
        )
    
    with col2:
        if st.button("📋 Copier dans le presse-papier", use_container_width=True):
            st.success("✅ Email copié ! (Ctrl+V pour coller)")

# ========================================
# PAGE : LIVRABLES WORD
# ========================================
elif page == "📄 Livrables Word":
    st.markdown(render_premium_header(
        title="📄 Livrables Word",
        subtitle="Génération rapports audit modifiables - Personnalisés par client"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        
        if not DOCX_AVAILABLE:
            st.error("❌ Le module python-docx n'est pas installé. Installez-le avec : `pip install python-docx`")
        else:
            st.markdown(section_divider("📋 SÉLECTION CLIENT"), unsafe_allow_html=True)
            
            # Sélection client
            client_names = df['NOM'].tolist()
            selected_client = st.selectbox("Choisissez un client", client_names)
            
            client_data = df[df['NOM'] == selected_client].iloc[0]
            
            # Aperçu données client
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**Informations client**")
                st.markdown(f"- **Nom :** {client_data['NOM']}")
                st.markdown(f"- **Secteur :** {client_data['SECTEUR']}")
                st.markdown(f"- **CA Honoraires :** {client_data['CA_HONORAIRES_HT']:,} €".replace(",", " "))
                st.markdown(f"- **Segment :** {client_data['SEGMENT']}")
            
            with col2:
                st.markdown("**Évaluation**")
                st.markdown(f"- **Outil :** {client_data['OUTIL_COMPATIBLE_REFORME']}")
                st.markdown(f"- **Appétence :** {client_data['APPETENCE_INFORMATIQUE']}")
                st.markdown(f"- **Score :** {client_data['SCORE_OPPORTUNITE']:.1f}")
                st.markdown(f"- **Priorité :** {client_data['PRIORITE']}")
            
            st.markdown(section_divider("⚙️ GÉNÉRATION RAPPORT"), unsafe_allow_html=True)
            
            if st.button("📄 Générer le rapport Word", use_container_width=True, type="primary"):
                try:
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)
                    
                    # ---------- RAPPORT P1 : AUDIT COMPLET (Priorité 1) ----------
                    if 'PRIORITÉ 1' in str(client_data.get('PRIORITE', '')):
                        p_version = doc.add_paragraph('VERSION 1 - AUDIT COMPLET (Priorité 1)')
                        p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p_version.runs:
                            p_version.runs[0].bold = True
                            p_version.runs[0].font.size = Pt(10)
                        doc.add_paragraph()
                        
                        header = doc.add_heading('RAPPORT D\'AUDIT', 0)
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_heading('Facturation Électronique 2026', level=2)
                        
                        p = doc.add_paragraph()
                        p.add_run('Client : ').bold = True
                        p.add_run(str(client_data['NOM']))
                        p = doc.add_paragraph()
                        p.add_run('Dirigeant : ').bold = True
                        dir_prenom = client_data.get('DIRIGEANT_PRENOM') or ''
                        dir_nom = client_data.get('DIRIGEANT_NOM') or ''
                        p.add_run(f"{dir_prenom} {dir_nom}".strip() or '[Prénom Nom]')
                        p = doc.add_paragraph()
                        p.add_run('Secteur d\'activité : ').bold = True
                        p.add_run(str(client_data.get('SECTEUR', '')))
                        p = doc.add_paragraph()
                        p.add_run('Date : ').bold = True
                        p.add_run(datetime.now().strftime('%d/%m/%Y'))
                        doc.add_paragraph()
                        
                        doc.add_heading('1. VOTRE SITUATION ACTUELLE', level=1)
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Outil de facturation :').bold = True
                        outil = str(client_data.get('OUTIL_COMPATIBLE_REFORME', '')).upper()
                        if outil == 'NON':
                            doc.add_paragraph(
                                'Votre solution actuelle n\'est pas conforme à la réforme. Un changement '
                                'd\'outil ou une migration vers une solution compatible est nécessaire. '
                                'Des paramétrages et une mise en relation avec une plateforme agréée (PDP/OD) '
                                'seront requis pour assurer la transmission des factures.'
                            )
                        else:
                            doc.add_paragraph(
                                'Votre solution actuelle nécessite une mise à jour auprès de votre éditeur '
                                'pour être en conformité avec la réforme. Des paramétrages complémentaires '
                                'seront également requis pour assurer la transmission des factures via une plateforme agréée.'
                            )
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Équipement numérique :').bold = True
                        appet = str(client_data.get('APPETENCE_INFORMATIQUE', '')).upper()
                        if appet in ('TRES BON', 'BON'):
                            doc.add_paragraph(
                                'Vous disposez d\'une maîtrise convenable des outils informatiques, ce qui facilitera '
                                'l\'appropriation des nouveaux process de facturation électronique.'
                            )
                        elif appet == 'MOYEN':
                            doc.add_paragraph(
                                'Votre niveau en outils informatiques est moyen. Un accompagnement et une formation '
                                'adaptée vous permettront de gagner en autonomie sur la facturation électronique.'
                            )
                        else:
                            doc.add_paragraph(
                                'Un accompagnement renforcé est recommandé pour vous familiariser avec les outils '
                                'et les nouveaux process de facturation électronique.'
                            )
                        doc.add_paragraph()
                        
                        doc.add_heading('2. LA RÉFORME DE LA FACTURATION ÉLECTRONIQUE', level=1)
                        doc.add_heading('2.1 - Échéance obligatoire', level=2)
                        doc.add_paragraph(
                            'À compter du 1er septembre 2026, toutes les entreprises assujetties à la TVA en France '
                            'devront basculer vers la facturation électronique pour leurs transactions B2B (entre professionnels).'
                        )
                        doc.add_paragraph()
                        doc.add_heading('2.2 - Vos obligations concrètes', level=2)
                        p = doc.add_paragraph()
                        p.add_run('ÉMISSION DE FACTURES').bold = True
                        doc.add_paragraph('Toutes vos factures clients devront être émises au format électronique structuré (et non plus en PDF simple).', style='List Bullet')
                        doc.add_paragraph('Le format privilégié est le format Factur-X (PDF lisible + données structurées XML).', style='List Bullet')
                        doc.add_paragraph('Les factures devront obligatoirement transiter par une plateforme de dématérialisation partenaire (PDP) ou un opérateur de dématérialisation (OD).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('RÉCEPTION DE FACTURES').bold = True
                        doc.add_paragraph('Vous devez être en mesure de recevoir et traiter des factures électroniques envoyées par vos fournisseurs.', style='List Bullet')
                        doc.add_paragraph('Votre système doit pouvoir lire les formats structurés (Factur-X, UBL, CII).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('TRANSMISSION DES DONNÉES').bold = True
                        doc.add_paragraph('Les données de facturation devront être transmises à l\'administration fiscale via le Portail Public de Facturation (PPF) ou via votre plateforme agréée.', style='List Bullet')
                        doc.add_paragraph('Cela inclut les données de transactions (e-reporting) et, pour certaines entreprises, les données de paiement (e-invoicing).', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('2.3 - Conséquences en cas de non-conformité', level=2)
                        doc.add_paragraph('Risque opérationnel : impossibilité d\'émettre ou de recevoir des factures conformes = blocage de votre activité commerciale.', style='List Bullet')
                        doc.add_paragraph('Risque fiscal : non-conformité = sanctions possibles et absence de déductibilité de la TVA sur les factures non conformes.', style='List Bullet')
                        doc.add_paragraph('Risque relationnel : vos clients et fournisseurs attendent de vous une conformité pour fluidifier les échanges.', style='List Bullet')
                        doc.add_paragraph()
                        
                        doc.add_heading('3. NOTRE ACCOMPAGNEMENT RECOMMANDÉ', level=1)
                        p = doc.add_paragraph()
                        p.add_run('Type de mission : ').bold = True
                        p.add_run('AUDIT COMPLET')
                        doc.add_paragraph()
                        doc.add_paragraph(
                            'Compte tenu de la nécessité de mise en conformité de votre outil, nous vous recommandons '
                            'un accompagnement complet structuré en 4 phases sur 8 semaines :'
                        )
                        doc.add_paragraph()
                        table_plan = doc.add_table(rows=5, cols=3)
                        table_plan.style = 'Table Grid'
                        table_plan.rows[0].cells[0].text = 'Phase'
                        table_plan.rows[0].cells[1].text = 'Durée'
                        table_plan.rows[0].cells[2].text = 'Contenu'
                        for j in range(3):
                            for run in table_plan.rows[0].cells[j].paragraphs[0].runs:
                                run.bold = True
                        table_plan.rows[1].cells[0].text = '1. Diagnostic approfondi'
                        table_plan.rows[1].cells[1].text = 'Semaines 1-2'
                        table_plan.rows[1].cells[2].text = 'Audit de votre outil actuel\nCartographie de vos flux de facturation (volume, typologie clients/fournisseurs)\nIdentification des prérequis techniques'
                        table_plan.rows[2].cells[0].text = '2. Mise en conformité'
                        table_plan.rows[2].cells[1].text = 'Semaines 3-5'
                        table_plan.rows[2].cells[2].text = 'Accompagnement au choix de la solution (mise à jour ou changement)\nParamétrage de l\'outil et de la plateforme de dématérialisation\nTests en environnement réel'
                        table_plan.rows[3].cells[0].text = '3. Formation'
                        table_plan.rows[3].cells[1].text = 'Semaine 6'
                        table_plan.rows[3].cells[2].text = 'Session de 2 heures avec vos équipes\nManipulation en conditions réelles\nCas pratiques adaptés à votre activité'
                        table_plan.rows[4].cells[0].text = '4. Suivi post-démarrage'
                        table_plan.rows[4].cells[1].text = 'Semaines 7-8'
                        table_plan.rows[4].cells[2].text = 'Points hebdomadaires de suivi\nAjustements et correction des anomalies\nValidation de la conformité'
                        doc.add_paragraph()
                        doc.add_paragraph('Tarif indicatif : 1 200 € à 1 500 € HT selon la complexité du dossier.')
                        doc.add_paragraph()
                        
                        doc.add_heading('4. PROCHAINES ÉTAPES', level=1)
                        doc.add_paragraph('Délai de mise en œuvre : 2 à 3 mois selon la complexité de votre situation.')
                        doc.add_paragraph()
                        doc.add_paragraph('Calendrier recommandé :')
                        doc.add_paragraph('Mars 2026 : Démarrage de l\'audit et choix de la solution', style='List Bullet')
                        doc.add_paragraph('Avril-Mai 2026 : Mise en place technique et formation', style='List Bullet')
                        doc.add_paragraph('Juin 2026 : Tests et ajustements finaux', style='List Bullet')
                        doc.add_paragraph('Septembre 2026 : Mise en production conforme', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_paragraph('Nous vous proposons un rendez-vous cette semaine ou la semaine prochaine pour :')
                        doc.add_paragraph('Valider ce diagnostic', style='List Bullet')
                        doc.add_paragraph('Préciser vos contraintes opérationnelles', style='List Bullet')
                        doc.add_paragraph('Planifier le démarrage de la mission', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_paragraph('_' * 60)
                        doc.add_paragraph()
                        footer_p = doc.add_paragraph()
                        footer_p.add_run('[Nom du cabinet]').bold = True
                        doc.add_paragraph('[Adresse]')
                        doc.add_paragraph('[Téléphone] | [Email] | [Site web]')
                    
                    # ---------- RAPPORT P2 : FORMATION (Priorité 2) ----------
                    elif 'PRIORITÉ 2' in str(client_data.get('PRIORITE', '')):
                        p_version = doc.add_paragraph('VERSION 2 - FORMATION (Priorité 2)')
                        p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p_version.runs:
                            p_version.runs[0].bold = True
                            p_version.runs[0].font.size = Pt(10)
                        doc.add_paragraph()
                        header = doc.add_heading('RAPPORT D\'AUDIT', 0)
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_heading('Facturation Électronique 2026', level=2)
                        p = doc.add_paragraph()
                        p.add_run('Client : ').bold = True
                        p.add_run(str(client_data['NOM']))
                        p = doc.add_paragraph()
                        p.add_run('Dirigeant : ').bold = True
                        dir_prenom = client_data.get('DIRIGEANT_PRENOM') or ''
                        dir_nom = client_data.get('DIRIGEANT_NOM') or ''
                        p.add_run(f"{dir_prenom} {dir_nom}".strip() or '[Prénom Nom]')
                        p = doc.add_paragraph()
                        p.add_run('Secteur d\'activité : ').bold = True
                        p.add_run(str(client_data.get('SECTEUR', '')))
                        p = doc.add_paragraph()
                        p.add_run('Date : ').bold = True
                        p.add_run(datetime.now().strftime('%d/%m/%Y'))
                        doc.add_paragraph()
                        doc.add_heading('1. VOTRE SITUATION ACTUELLE', level=1)
                        p = doc.add_paragraph()
                        p.add_run('Outil de facturation :').bold = True
                        doc.add_paragraph(
                            'Votre solution actuelle est compatible avec la réforme de la facturation électronique. '
                            'Votre éditeur a prévu les mises à jour nécessaires pour assurer la conformité au 1er septembre 2026.'
                        )
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Équipement numérique :').bold = True
                        doc.add_paragraph(
                            'Vous disposez d\'une bonne maîtrise des outils informatiques, ce qui vous permettra '
                            'd\'adopter rapidement les nouveaux process de facturation électronique.'
                        )
                        doc.add_paragraph()
                        doc.add_heading('2. LA RÉFORME DE LA FACTURATION ÉLECTRONIQUE', level=1)
                        doc.add_heading('2.1 - Échéance obligatoire', level=2)
                        doc.add_paragraph(
                            'À compter du 1er septembre 2026, toutes les entreprises assujetties à la TVA en France '
                            'devront basculer vers la facturation électronique pour leurs transactions B2B (entre professionnels).'
                        )
                        doc.add_paragraph()
                        doc.add_heading('2.2 - Vos obligations concrètes', level=2)
                        p = doc.add_paragraph()
                        p.add_run('ÉMISSION DE FACTURES').bold = True
                        doc.add_paragraph('Toutes vos factures clients devront être émises au format électronique structuré (et non plus en PDF simple).', style='List Bullet')
                        doc.add_paragraph('Le format privilégié est le format Factur-X (PDF lisible + données structurées XML).', style='List Bullet')
                        doc.add_paragraph('Les factures devront obligatoirement transiter par une plateforme de dématérialisation partenaire (PDP) ou un opérateur de dématérialisation (OD).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('RÉCEPTION DE FACTURES').bold = True
                        doc.add_paragraph('Vous devez être en mesure de recevoir et traiter des factures électroniques envoyées par vos fournisseurs.', style='List Bullet')
                        doc.add_paragraph('Votre système doit pouvoir lire les formats structurés (Factur-X, UBL, CII).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('TRANSMISSION DES DONNÉES').bold = True
                        doc.add_paragraph('Les données de facturation devront être transmises à l\'administration fiscale via le Portail Public de Facturation (PPF) ou via votre plateforme agréée.', style='List Bullet')
                        doc.add_paragraph('Cela inclut les données de transactions (e-reporting) et, pour certaines entreprises, les données de paiement (e-invoicing).', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('2.3 - Points de vigilance', level=2)
                        doc.add_paragraph('Mentions obligatoires : Les factures électroniques doivent contenir des mentions spécifiques (numéro SIREN, adresse complète, statut TVA, etc.)', style='List Bullet')
                        doc.add_paragraph('Archivage : Les factures électroniques doivent être conservées sous format électronique pendant 10 ans.', style='List Bullet')
                        doc.add_paragraph('Gestion des rejets : Une facture non conforme sera automatiquement rejetée par la plateforme — il est crucial de maîtriser les règles de contrôle.', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('3. NOTRE ACCOMPAGNEMENT RECOMMANDÉ', level=1)
                        p = doc.add_paragraph()
                        p.add_run('Type de mission : ').bold = True
                        p.add_run('FORMATION OPÉRATIONNELLE')
                        doc.add_paragraph()
                        doc.add_paragraph(
                            'Votre outil étant conforme, l\'enjeu est maintenant de rendre vos équipes autonomes '
                            'pour éviter les erreurs et les rejets de factures.'
                        )
                        doc.add_paragraph()
                        doc.add_paragraph('Session de formation pratique - 3 heures')
                        doc.add_paragraph()
                        table_form = doc.add_table(rows=4, cols=2)
                        table_form.style = 'Table Grid'
                        table_form.rows[0].cells[0].text = 'Séquence'
                        table_form.rows[0].cells[1].text = 'Contenu'
                        for j in range(2):
                            for run in table_form.rows[0].cells[j].paragraphs[0].runs:
                                run.bold = True
                        table_form.rows[1].cells[0].text = '1. Comprendre la réforme (sans jargon)'
                        table_form.rows[1].cells[1].text = 'Obligations légales simplifiées\nImpacts concrets sur votre activité\nCalendrier et étapes clés'
                        table_form.rows[2].cells[0].text = '2. Manipuler votre outil en conditions réelles'
                        table_form.rows[2].cells[1].text = 'Paramétrage de base\nÉmission d\'une facture conforme\nRéception et traitement d\'une facture électronique\nGestion des anomalies'
                        table_form.rows[3].cells[0].text = '3. Éviter les erreurs classiques'
                        table_form.rows[3].cells[1].text = 'Les 5 erreurs qui bloquent une facture\nProcédure en cas de rejet\nQui contacter en cas de problème'
                        doc.add_paragraph()
                        doc.add_paragraph('Livrable inclus : Support de formation complet à conserver.')
                        doc.add_paragraph()
                        doc.add_heading('4. PROCHAINES ÉTAPES', level=1)
                        doc.add_paragraph('Nous vous proposons un rendez-vous cette semaine ou la semaine prochaine pour :')
                        doc.add_paragraph('Confirmer vos besoins', style='List Bullet')
                        doc.add_paragraph('Planifier la session de formation', style='List Bullet')
                        doc.add_paragraph('Répondre à vos questions', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_paragraph('_' * 60)
                        doc.add_paragraph()
                        footer_p = doc.add_paragraph()
                        footer_p.add_run('[Nom du cabinet]').bold = True
                        doc.add_paragraph('[Adresse]')
                        doc.add_paragraph('[Téléphone] | [Email] | [Site web]')
                    
                    # ---------- RAPPORT P3 : INFORMATION (Priorité 3) ----------
                    else:
                        p_version = doc.add_paragraph('VERSION 3 - INFORMATION (Priorité 3)')
                        p_version.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p_version.runs:
                            p_version.runs[0].bold = True
                            p_version.runs[0].font.size = Pt(10)
                        doc.add_paragraph()
                        header = doc.add_heading('RAPPORT D\'AUDIT', 0)
                        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_heading('Facturation Électronique 2026', level=2)
                        p = doc.add_paragraph()
                        p.add_run('Client : ').bold = True
                        p.add_run(str(client_data['NOM']))
                        p = doc.add_paragraph()
                        p.add_run('Dirigeant : ').bold = True
                        dir_prenom = client_data.get('DIRIGEANT_PRENOM') or ''
                        dir_nom = client_data.get('DIRIGEANT_NOM') or ''
                        p.add_run(f"{dir_prenom} {dir_nom}".strip() or '[Prénom Nom]')
                        p = doc.add_paragraph()
                        p.add_run('Secteur d\'activité : ').bold = True
                        p.add_run(str(client_data.get('SECTEUR', '')))
                        p = doc.add_paragraph()
                        p.add_run('Date : ').bold = True
                        p.add_run(datetime.now().strftime('%d/%m/%Y'))
                        doc.add_paragraph()
                        doc.add_heading('1. VOTRE SITUATION ACTUELLE', level=1)
                        p = doc.add_paragraph()
                        p.add_run('Outil de facturation :').bold = True
                        doc.add_paragraph(
                            'Votre solution actuelle est pleinement compatible avec la réforme de la facturation électronique. '
                            'Votre éditeur a confirmé que les mises à jour nécessaires seront déployées automatiquement avant le 1er septembre 2026.'
                        )
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Équipement numérique :').bold = True
                        doc.add_paragraph(
                            'Vous disposez d\'une très bonne maîtrise des outils informatiques et d\'une bonne organisation administrative, '
                            'ce qui vous place dans une situation favorable pour aborder cette transition sereinement.'
                        )
                        doc.add_paragraph()
                        doc.add_heading('2. LA RÉFORME DE LA FACTURATION ÉLECTRONIQUE', level=1)
                        doc.add_heading('2.1 - Échéance obligatoire', level=2)
                        doc.add_paragraph(
                            'À compter du 1er septembre 2026, toutes les entreprises assujetties à la TVA en France '
                            'devront basculer vers la facturation électronique pour leurs transactions B2B (entre professionnels).'
                        )
                        doc.add_paragraph()
                        doc.add_heading('2.2 - Vos obligations concrètes', level=2)
                        p = doc.add_paragraph()
                        p.add_run('ÉMISSION DE FACTURES').bold = True
                        doc.add_paragraph('Toutes vos factures clients devront être émises au format électronique structuré (et non plus en PDF simple).', style='List Bullet')
                        doc.add_paragraph('Le format privilégié est le format Factur-X (PDF lisible + données structurées XML).', style='List Bullet')
                        doc.add_paragraph('Les factures devront obligatoirement transiter par une plateforme de dématérialisation partenaire (PDP) ou un opérateur de dématérialisation (OD).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('RÉCEPTION DE FACTURES').bold = True
                        doc.add_paragraph('Vous devez être en mesure de recevoir et traiter des factures électroniques envoyées par vos fournisseurs.', style='List Bullet')
                        doc.add_paragraph('Votre système doit pouvoir lire les formats structurés (Factur-X, UBL, CII).', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('TRANSMISSION DES DONNÉES').bold = True
                        doc.add_paragraph('Les données de facturation devront être transmises à l\'administration fiscale via le Portail Public de Facturation (PPF) ou via votre plateforme agréée.', style='List Bullet')
                        doc.add_paragraph('Cela inclut les données de transactions (e-reporting) et, pour certaines entreprises, les données de paiement (e-invoicing).', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('2.3 - Les grands principes', level=2)
                        doc.add_paragraph('Authenticité : L\'origine de la facture doit être garantie (signature électronique ou piste d\'audit fiable).', style='List Bullet')
                        doc.add_paragraph('Intégrité : Le contenu ne peut pas être modifié après émission.', style='List Bullet')
                        doc.add_paragraph('Lisibilité : La facture doit être lisible par l\'humain ET par les systèmes informatiques.', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('3. POINTS DE VIGILANCE (SIMPLES)', level=1)
                        doc.add_paragraph(
                            'Bien que vous soyez en bonne position, voici les quelques vérifications à effectuer d\'ici septembre 2026 :'
                        )
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Vérifier vos coordonnées administratives').bold = True
                        doc.add_paragraph('SIRET à jour', style='List Bullet')
                        doc.add_paragraph('Adresse complète et exacte', style='List Bullet')
                        doc.add_paragraph('Email de facturation fonctionnel', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Suivre les mises à jour de votre logiciel').bold = True
                        doc.add_paragraph('Installer les mises à jour proposées par votre éditeur (automne 2026)', style='List Bullet')
                        doc.add_paragraph('Vérifier que votre abonnement inclut bien la conformité facturation électronique', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Tester l\'envoi d\'une première facture électronique').bold = True
                        doc.add_paragraph('Effectuer un test en conditions réelles (septembre 2026)', style='List Bullet')
                        doc.add_paragraph('Vérifier la réception et le traitement par votre client', style='List Bullet')
                        doc.add_paragraph()
                        p = doc.add_paragraph()
                        p.add_run('Informer vos principaux clients et fournisseurs').bold = True
                        doc.add_paragraph('Les sensibiliser à la réforme', style='List Bullet')
                        doc.add_paragraph('Confirmer leurs coordonnées de facturation électronique', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('4. NOTRE RECOMMANDATION', level=1)
                        p = doc.add_paragraph()
                        p.add_run('Validation de conformité - Point téléphonique de 30 minutes').bold = True
                        doc.add_paragraph()
                        doc.add_paragraph(
                            'Bien que votre situation soit favorable, nous vous recommandons un point de validation pour :'
                        )
                        doc.add_paragraph('Confirmer que tous les voyants sont au vert', style='List Bullet')
                        doc.add_paragraph('Répondre à vos éventuelles questions', style='List Bullet')
                        doc.add_paragraph('Vous remettre une checklist finale avant septembre 2026', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_paragraph('Calendrier recommandé :')
                        doc.add_paragraph('Mars-Juin 2026 : Vérifications administratives', style='List Bullet')
                        doc.add_paragraph('Juillet-Août 2026 : Installation des mises à jour logiciel', style='List Bullet')
                        doc.add_paragraph('Septembre 2026 : Premier envoi test + mise en production', style='List Bullet')
                        doc.add_paragraph()
                        doc.add_heading('5. PROCHAINES ÉTAPES', level=1)
                        doc.add_paragraph(
                            'Nous vous proposons un point téléphonique cette semaine ou la semaine prochaine pour valider votre situation.'
                        )
                        doc.add_paragraph('Ce point est inclus dans nos honoraires habituels.')
                        doc.add_paragraph()
                        doc.add_paragraph('_' * 60)
                        doc.add_paragraph()
                        footer_p = doc.add_paragraph()
                        footer_p.add_run('[Nom du cabinet]').bold = True
                        doc.add_paragraph('[Adresse]')
                        doc.add_paragraph('[Téléphone] | [Email] | [Site web]')
                    
                    output_docx = io.BytesIO()
                    doc.save(output_docx)
                    output_docx.seek(0)
                    st.success("✅ Rapport généré avec succès !")
                    st.download_button(
                        label="📥 Télécharger le rapport Word",
                        data=output_docx.getvalue(),
                        file_name=f"rapport_audit_{client_data['NOM'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"❌ Erreur lors de la génération : {str(e)}")

# ========================================
# PAGE : BUDGET TCO
# ========================================
elif page == "🔧 Budget TCO":
    st.markdown(render_premium_header(
        title="🔧 Comparateur Budget TCO",
        subtitle="Total Cost of Ownership — Solutions facturation électronique 2026"
    ), unsafe_allow_html=True)
    
    with st.expander("ℹ️ Qu’est-ce que le TCO et comment lire ce comparateur ?", expanded=True):
        st.markdown("""
        **TCO (Total Cost of Ownership)** = coût total de possession sur la période choisie.
        
        - **Coût initial** : mise en place (formation, paramétrage, éventuel setup).
        - **Coût mensuel** : abonnement (forfait + utilisateurs). La transmission entre Plateformes Agréées (PA) est gratuite — pas de coût par facture pour l'échange inter-PA.
        - **TCO** = Coût initial + (Coût mensuel × nombre de mois).
        
        Les montants sont **indicatifs** et basés sur les grilles tarifaires publiées ou des fourchettes courantes.  
        Les offres évoluent : **vérifiez les tarifs à jour sur les sites des éditeurs** avant toute décision.
        """)
    
    st.markdown(section_divider("⚙️ PARAMÈTRES CLIENT"), unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        nb_users = st.number_input("Nombre d'utilisateurs", 1, 50, 3, help="Utilisateurs qui émettent ou traitent des factures")
    
    with col2:
        nb_factures = st.number_input("Factures / mois", 10, 2000, 100, help="Volume moyen de factures émises ou reçues par mois")
    
    with col3:
        horizon = st.selectbox("Horizon (mois)", [12, 24, 36], index=1, help="Période sur laquelle est calculé le TCO")
    
    st.markdown(section_divider("💰 COMPARAISON SOLUTIONS"), unsafe_allow_html=True)
    
    # Données solutions — abonnements uniquement (pas de coût par facture : interopérabilité entre PA gratuite)
    solutions = {
        "Pennylane": {
            "prix_base": 49,
            "prix_user": 15,
            "cout_formation": 400,
            "delai_deploiement": "2 semaines",
            "support": "Email + Chat",
            "source": "pennylane.com — Plan Basique 1-5 sal., e-facture incluse"
        },
        "Cegid Loop": {
            "prix_base": 89,
            "prix_user": 15,
            "cout_formation": 800,
            "delai_deploiement": "4 semaines",
            "support": "Téléphone + Email",
            "source": "Indicatif — Cegid Loop : tarifs sur devis (lebonlogiciel.com, cegid.com)"
        },
        "Sage 100 (Sage Network PA)": {
            "prix_base": 120,
            "prix_user": 20,
            "cout_formation": 1000,
            "delai_deploiement": "6 semaines",
            "support": "Premium",
            "source": "Indicatif — Sage : tarifs sur devis (blc-conseil.com)"
        },
        "Inqom": {
            "prix_base": 199,
            "prix_user": 0,
            "cout_formation": 0,
            "delai_deploiement": "2 mois (lancement)",
            "support": "Email + Tél. 01 84 80 25 56",
            "source": "inqom.com/tarifs — 199 €/mois (licences + accompagnement sur mesure). Facture électronique incluse."
        },
        "Tiime AE": {
            "prix_base": 59,
            "prix_user": 12,
            "cout_formation": 500,
            "delai_deploiement": "3 semaines",
            "support": "Chat + Email",
            "source": "tiime.fr — Tarifs cabinet sur devis / démo. PA facturation électronique."
        }
    }
    
    # Calculs — coût mensuel = abonnement (base + utilisateurs) uniquement ; pas de coût par facture (transmission entre PA gratuite)
    results = []
    for nom, params in solutions.items():
        cout_mensuel = params['prix_base'] + (nb_users * params['prix_user'])
        cout_initial = params['cout_formation']
        tco = (cout_mensuel * horizon) + cout_initial
        results.append({
            "Solution": nom,
            "Coût mensuel": f"{cout_mensuel:.2f} €",
            "Coût initial": f"{cout_initial} €",
            "TCO ({} mois)".format(horizon): f"{tco:.2f} €",
            "Délai": params['delai_deploiement'],
            "Support": params['support'],
            "Source": params['source'],
            "_tco_num": tco
        })
    
    # Tri par TCO croissant — Pennylane apparaît bien placé quand son TCO est compétitif
    results_sorted = sorted(results, key=lambda r: r["_tco_num"])
    for r in results_sorted:
        r.pop("_tco_num", None)
    df_results = pd.DataFrame(results_sorted)
    
    st.dataframe(df_results, use_container_width=True, hide_index=True, column_config={"Source": st.column_config.TextColumn("Source / remarque", width="large")})
    
    st.caption("Les montants sont indicatifs. Inqom : inqom.com/tarifs (199 €/mois). Tiime : tiime.fr — tarifs cabinet sur devis. Vérifiez les grilles à jour sur les sites des éditeurs.")
    
    # Graphique
    st.markdown(section_divider("📊 VISUALISATION TCO"), unsafe_allow_html=True)
    
    tco_values = [float(r["TCO ({} mois)".format(horizon)].replace(" €", "").replace(",", ".")) for r in results_sorted]
    solution_names = [r["Solution"] for r in results_sorted]
    colors = ['#10b981', '#3b82f6', '#f59e0b', '#ef4444', '#8b5cf6', '#06b6d4']
    
    fig_tco = go.Figure(data=[
        go.Bar(
            x=solution_names,
            y=tco_values,
            text=[f"{v:.0f} €" for v in tco_values],
            textposition='outside',
            marker=dict(color=colors[:len(solution_names)])
        )
    ])
    
    fig_tco.update_layout(
        height=400,
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(15,23,42,0.8)',
        font=dict(color='#f1f5f9'),
        xaxis=dict(title="", gridcolor='rgba(255,255,255,0.1)', tickangle=-25),
        yaxis=dict(title=f"TCO sur {horizon} mois (€)", gridcolor='rgba(255,255,255,0.1)')
    )
    
    st.plotly_chart(fig_tco, use_container_width=True)

# ========================================
# PAGE : EXPORTS
# ========================================
elif page == "📤 Exports":
    st.markdown(render_premium_header(
        title="📤 Exports de Données",
        subtitle="Téléchargement Excel enrichi avec toutes les analyses"
    ), unsafe_allow_html=True)
    
    if st.session_state.df is None:
        st.warning("⚠️ Veuillez d'abord importer vos données depuis la page Accueil")
    else:
        df = st.session_state.df
        
        st.markdown(section_divider("📊 DONNÉES DISPONIBLES"), unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("Nombre de clients", len(df))
            st.metric("Colonnes", len(df.columns))
        
        with col2:
            file_size_mb = len(df.to_csv().encode('utf-8')) / (1024 * 1024)
            st.metric("Taille estimée", f"{file_size_mb:.2f} Mo")
            st.metric("Format", "Excel (.xlsx)")
        
        st.markdown(section_divider("📥 TÉLÉCHARGEMENT"), unsafe_allow_html=True)
        
        st.markdown("""
        **Ce fichier contient :**
        - ✅ Toutes les colonnes d'origine (NOM, SECTEUR, CA, etc.)
        - ✅ Toutes les colonnes calculées (SEGMENT, SCORE, PRIORITÉ, ETOILES)
        - ✅ Format Excel (.xlsx) compatible avec tous les outils
        """)
        
        # Génération Excel
        output_export = io.BytesIO()
        
        with pd.ExcelWriter(output_export, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name='Analyse complète', index=False)
            
            workbook = writer.book
            worksheet = writer.sheets['Analyse complète']
            
            # Format header
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'vcenter',
                'fg_color': '#667eea',
                'font_color': 'white',
                'border': 1
            })
            
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(0, col_num, value, header_format)
            
            # Largeurs colonnes
            worksheet.set_column('A:A', 25)  # NOM
            worksheet.set_column('B:B', 20)  # SECTEUR
            worksheet.set_column('C:C', 15)  # CA
            worksheet.set_column('D:E', 25)  # OUTIL, APPETENCE
            worksheet.set_column('F:H', 15)  # Dirigeant
            worksheet.set_column('I:L', 20)  # Calculées
        
        output_export.seek(0)
        
        st.download_button(
            label="📥 Télécharger l'analyse complète (Excel)",
            data=output_export.getvalue(),
            file_name=f"analyse_fae_complete_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            type="primary"
        )
        
        st.markdown("---")
        
        st.info("""
        💡 **Astuce** : Ce fichier peut être réimporté dans l'outil pour conserver vos analyses. 
        Les colonnes calculées seront écrasées lors d'un nouvel import.
        """)

# ========================================
# FOOTER GLOBAL
# ========================================
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem 0; color: #64748b; font-size: 0.9rem;">
    <p style="margin: 0;">🏆 <strong>Outil d'Analyse & Segmentation Client RFE v7.3</strong></p>
    <p style="margin: 0.5rem 0 0 0;">Développé avec ❤️ par Jesse DEVENON pour les cabinets d'expertise comptable</p>
    <p style="margin: 0.5rem 0 0 0;">🔒 Traitement 100% local - Confidentialité garantie</p>
</div>
""", unsafe_allow_html=True)